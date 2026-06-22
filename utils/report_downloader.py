import streamlit as st
import requests
import datetime
import io
import math
from utils.config import API_BASE_URL, IS_API_MODE


# ═══════════════════════════════════════════════════════════════════════════════
# Expert Domain Knowledge — derived from 02_생활정보/00_Skill_MD/life-*.md (이전됨 2026-06-18)
# ═══════════════════════════════════════════════════════════════════════════════

EXPERT_ANALYSIS = {
    "주식": {
        "icon": "📈",
        "framework": "기술적 분석 + 펀더멘털 분석",
        "metrics": ["PER", "PBR", "ROE", "EPS", "배당수익률"],
        "disclaimer": "본 분석은 투자 참고용이며, 투자 판단의 최종 책임은 투자자 본인에게 있습니다. 과거 수익률이 미래 수익률을 보장하지 않습니다.",
        "analysis_fn": "_analyze_stock",
    },
    "환율": {
        "icon": "💱",
        "framework": "거시경제 환율 분석",
        "metrics": ["USD/KRW", "JPY/KRW", "CNY/KRW", "환율변동성"],
        "disclaimer": "환율은 글로벌 거시경제, 금리 정책, 지정학적 리스크 등 복합 요인에 의해 변동됩니다.",
        "analysis_fn": "_analyze_forex",
    },
    "관세": {
        "icon": "🚢",
        "framework": "무역수지 및 관세영향 분석",
        "metrics": ["수출증감률", "수입증감률", "무역수지", "관세율변동"],
        "disclaimer": "관세 및 무역 정책은 국제 협정과 정부 정책에 따라 변동될 수 있습니다.",
        "analysis_fn": "_analyze_trade",
    },
    "금융": {
        "icon": "💰",
        "framework": "가계금융 및 자산관리 분석",
        "metrics": ["기준금리", "예금금리", "대출금리", "소비자물가지수"],
        "disclaimer": "금융 상품 비교 시 수수료, 세금, 중도해지 조건을 반드시 확인하세요.",
        "analysis_fn": "_analyze_finance",
    },
    "건강": {
        "icon": "🏥",
        "framework": "헬스케어 트렌드 및 건강관리 분석",
        "metrics": ["BMI 기준", "일일권장섭취량", "운동강도지표"],
        "disclaimer": "본 분석은 전문 의료 상담을 대체하지 않습니다. 개인 건강 상태에 따라 전문의 상담을 권장합니다.",
        "analysis_fn": "_analyze_health",
    },
    "부동산": {
        "icon": "🏠",
        "framework": "부동산 시장 동향 및 가격 분석",
        "metrics": ["매매가격지수", "전세가격지수", "매매/전세비율", "청약경쟁률"],
        "disclaimer": "부동산 투자는 가격하락, 금리변동, 공실률 등의 위험을 수반합니다.",
        "analysis_fn": "_analyze_realestate",
    },
    "법률": {
        "icon": "⚖️",
        "framework": "생활법률 및 판례 분석",
        "metrics": ["관련법령", "판례동향", "소멸시효"],
        "disclaimer": "본 분석은 전문 법률 상담을 대체하지 않습니다. 대한법률구조공단(132) 무료 상담을 권장합니다.",
        "analysis_fn": "_analyze_legal",
    },
    "교육": {
        "icon": "📚",
        "framework": "에듀테크 및 학습 트렌드 분석",
        "metrics": ["수강률", "합격률", "학습효율지표"],
        "disclaimer": "",
        "analysis_fn": "_analyze_default",
    },
    "여행": {
        "icon": "✈️",
        "framework": "여행 트렌드 및 관광 분석",
        "metrics": ["항공운임지수", "관광객수", "숙박가격지수"],
        "disclaimer": "여행 전 외교부 해외안전여행 사이트에서 여행경보 단계를 확인하세요.",
        "analysis_fn": "_analyze_default",
    },
    "식생활": {
        "icon": "🍽️",
        "framework": "외식·식품 트렌드 분석",
        "metrics": ["외식물가지수", "식재료가격", "배달주문트렌드"],
        "disclaimer": "",
        "analysis_fn": "_analyze_default",
    },
    "쇼핑": {
        "icon": "🛒",
        "framework": "소비 트렌드 및 쇼핑 분석",
        "metrics": ["소비자심리지수", "온라인거래액", "카드소비동향"],
        "disclaimer": "",
        "analysis_fn": "_analyze_default",
    },
    "육아": {
        "icon": "👶",
        "framework": "육아·보육 정책 및 트렌드 분석",
        "metrics": ["합계출산율", "보육시설현황", "육아용품시장"],
        "disclaimer": "",
        "analysis_fn": "_analyze_default",
    },
    "문화": {
        "icon": "🎭",
        "framework": "문화·예술 동향 분석",
        "metrics": ["공연관람률", "전시회참관객", "한류콘텐츠수출"],
        "disclaimer": "",
        "analysis_fn": "_analyze_default",
    },
    "반려동물": {
        "icon": "🐾",
        "framework": "펫 산업 트렌드 분석",
        "metrics": ["반려동물등록수", "펫산업시장규모", "사료가격지수"],
        "disclaimer": "",
        "analysis_fn": "_analyze_default",
    },
    "화훼": {
        "icon": "🌸",
        "framework": "화훼·식물 트렌드 분석",
        "metrics": ["화훼시장규모", "플랜테리어관심도"],
        "disclaimer": "",
        "analysis_fn": "_analyze_default",
    },
    "유가": {
        "icon": "🛢️",
        "framework": "국제유가 분석",
        "metrics": ["WTI", "Brent", "Dubai유", "국내유가"],
        "disclaimer": "국제유가는 OPEC+ 감산, 지정학적 긴장, 글로벌 수요 등 복합 요인에 의해 변동됩니다.",
        "analysis_fn": "_analyze_default",
    },
    "운송": {
        "icon": "🚛",
        "framework": "물류·운송 동향 분석",
        "metrics": ["BDI(발틱건화물지수)", "컨테이너운임", "항공화물량"],
        "disclaimer": "",
        "analysis_fn": "_analyze_default",
    },
    "분쟁": {
        "icon": "🌍",
        "framework": "글로벌 지정학 리스크 분석",
        "metrics": ["지정학리스크지수", "분쟁영향권국가", "난민현황"],
        "disclaimer": "",
        "analysis_fn": "_analyze_default",
    },
    "사업": {
        "icon": "🏢",
        "framework": "창업·비즈니스 트렌드 분석",
        "metrics": ["창업건수", "폐업률", "VC투자규모"],
        "disclaimer": "",
        "analysis_fn": "_analyze_default",
    },
}


# 분야 식별 키워드(별칭) — 검색어에 분야명("주식"·"법률")이 없어도 정확히 매칭.
# 검색 종류에 맞는 전문 분석을 보장하는 핵심. 모두 소문자(combined가 lower 처리됨).
DOMAIN_ALIASES = {
    "주식": ["주식", "코스피", "코스닥", "증시", "종목", "주가", "나스닥", "다우", "s&p", "sp500",
            "per", "pbr", "roe", "etf", "상장", "배당", "매수", "매도", "반도체", "2차전지", "시황"],
    "환율": ["환율", "원달러", "원/달러", "달러인덱스", "dxy", "외환", "엔화", "위안", "금리차"],
    "관세": ["관세", "무역", "수출", "수입", "무역수지", "fta", "통상", "hs코드", "보호무역"],
    "금융": ["금융", "금리", "기준금리", "대출", "예금", "적금", "펀드", "보험", "카드", "가계부채",
            "자산관리", "통화정책", "재테크"],
    "건강": ["건강", "질병", "운동", "영양", "다이어트", "의료", "병원", "질환", "면역", "수면",
            "스트레스", "bmi", "헬스"],
    "부동산": ["부동산", "아파트", "전세", "매매", "청약", "분양", "주택", "임대", "ltv", "dsr",
             "재건축", "집값", "월세"],
    "법률": ["법률", "판례", "소송", "계약", "임대차", "보증금", "상속", "이혼", "손해배상", "형사",
            "민사", "고소", "고발", "소멸시효", "법령", "변호사"],
    "교육": ["교육", "학습", "입시", "수능", "학원", "강의", "자격증", "유학", "에듀", "공부", "시험"],
    "여행": ["여행", "항공", "숙박", "호텔", "관광", "비자", "여권", "패키지", "항공권", "해외여행"],
    "식생활": ["식생활", "외식", "식품", "음식", "맛집", "식재료", "배달", "레시피", "식단", "요리"],
    "쇼핑": ["쇼핑", "소비", "할인", "구매", "이커머스", "온라인쇼핑", "직구", "세일", "소비자심리"],
    "육아": ["육아", "출산", "보육", "어린이집", "유치원", "아동", "출산율", "양육", "분유", "기저귀"],
    "문화": ["문화", "공연", "전시", "영화", "콘서트", "뮤지컬", "한류", "예술", "박물관", "ott"],
    "반려동물": ["반려동물", "반려견", "반려묘", "강아지", "고양이", "사료", "동물병원", "펫"],
    "화훼": ["화훼", "플랜테리어", "화분", "원예", "가드닝", "절화"],
    "유가": ["유가", "wti", "brent", "브렌트", "두바이유", "원유", "휘발유", "경유", "opec", "정유"],
    "운송": ["운송", "물류", "해운", "운임", "bdi", "컨테이너", "화물", "항만", "택배", "scfi"],
    "분쟁": ["분쟁", "전쟁", "지정학", "군사", "난민", "제재", "무력", "안보", "우크라", "중동", "휴전"],
    "사업": ["사업", "창업", "스타트업", "벤처", "비즈니스", "폐업", "vc", "투자유치", "법인"],
}


def _match_expert_domain(query: str, expert_name: str = "") -> dict:
    """Match query/expert_name to the best expert domain analysis profile.
    분야명 직접 포함 + 분야 키워드(별칭) 가중 매칭 — 검색 종류에 맞는 전문 분석 보장."""
    combined = f"{expert_name} {query}".lower()
    best_match = None
    best_key = None
    best_score = 0
    for key, profile in EXPERT_ANALYSIS.items():
        key_lower = key.lower()
        score = 0
        if key_lower in combined:
            score += len(key_lower) * 3
        for alias in DOMAIN_ALIASES.get(key, ()):
            if alias and alias in combined:
                score += len(alias) * 3 + 2
        if score > best_score:
            best_score = score
            best_match = profile
            best_key = key
    if best_match and best_score >= 2:
        return {**best_match, "_key": best_key}
    # Default
    return {
        "icon": "📊",
        "framework": "데이터 기반 트렌드 분석",
        "metrics": ["관심도지수", "뉴스빈도", "검색트렌드"],
        "disclaimer": "",
        "analysis_fn": "_analyze_default",
        "_key": "default",
    }


# ═══════════════════════════════════════════════════════════════════════════════
# 분야별 전문 심층 분석 지식 (Domain Deep-Dive)
# 근거: 02_생활정보/00_Skill_MD/life-*.md(주식·금융·건강·부동산·법률·교육·여행·식생활)
#       + 각 분야 정통 분석 프레임워크. 죽어있던 analysis_fn(_analyze_*)을 실제 구현으로 대체.
# 각 분야: lens(전문가 관점) · indicators(핵심지표 해석) · risks · strategy · checklist
# ═══════════════════════════════════════════════════════════════════════════════

DOMAIN_DEEPDIVE = {
    "주식": {
        "lens": [
            ("밸류에이션", "PER·PBR을 업종 평균·과거 밴드와 비교 — PBR 1배 미만은 순자산 이하 거래, 고PER은 성장 선반영"),
            ("실적 모멘텀", "최근 8분기 매출·영업이익 추이와 YoY/QoQ 성장률, 컨센서스 대비 어닝 서프라이즈/미스"),
            ("수급 동향", "외국인·기관 순매수/순매도, 거래대금, 프로그램 매매 방향"),
            ("기술적 위치", "5/20/60/120일 이동평균 배열, RSI(과매수 70/과매도 30), MACD 교차, 지지·저항"),
            ("거시 연동", "기준금리·원/달러 환율·미 국채 10년물·반도체 사이클과의 연동성"),
        ],
        "indicators": [
            ("PER", "업종 평균 대비 낮으면 저평가, 높으면 성장 기대 선반영"),
            ("PBR", "1배 미만은 청산가치 이하 — 자산주 저평가 신호"),
            ("ROE", "15% 이상이면 자본 효율 우량"),
            ("RSI(14)", "70 이상 과매수 경계 / 30 이하 과매도 반등 가능"),
            ("베타(β)", "1 초과는 시장보다 변동 큼(공격적) / 1 미만 방어적"),
        ],
        "risks": ["시장 리스크(금리·환율)", "업종 리스크(반도체·2차전지 사이클)", "개별 리스크(실적쇼크·공시)", "수급 리스크(외국인 이탈)"],
        "strategy": ["가치투자: 저PER/PBR 저평가주 장기 보유", "성장투자: 고성장 매출·이익 종목 선별", "모멘텀: 강한 상승 추세 편승(단기)", "배당투자: 안정적 배당수익(장기·보수)"],
        "checklist": ["관심 종목 PER/PBR을 업종 평균과 비교했는가", "최근 분기 실적·가이던스를 확인했는가", "외국인·기관 수급 방향을 점검했는가", "진입·목표·손절 가격 시나리오를 설정했는가"],
        "source": "life-stock-analyst.md",
    },
    "환율": {
        "lens": [
            ("금리차·캐리", "내외 금리차 확대 시 자본이동·캐리 트레이드 방향 변화"),
            ("달러인덱스(DXY)", "달러 강세 국면에서는 신흥국·원화 동반 약세 경향"),
            ("대외건전성", "경상수지·외환보유고·자본수지로 통화 방어력 판단"),
            ("당국 개입·변동성", "급변동 구간의 스무딩 오퍼레이션·구두개입 경계"),
            ("글로벌 이벤트", "FOMC·지정학 리스크·위험선호(Risk-on/off) 전이"),
        ],
        "indicators": [
            ("원/달러", "상승=원화 약세 → 수출주 수혜·수입물가 상승"),
            ("DXY", "상승=달러 강세 → 신흥국 통화 약세 압력"),
            ("내외 금리차", "미국>한국 확대 시 자본유출·원화 약세 압력"),
            ("변동성", "일중 변동 확대 시 단기 투기·당국 개입 가능성"),
        ],
        "risks": ["미 통화정책 전환", "지정학 리스크", "경상수지 악화", "급격한 자본유출"],
        "strategy": ["분할 환전·환헤지로 변동성 분산", "수출입 네고 타이밍 최적화", "달러·통화 분산 보유"],
        "checklist": ["내외 금리차 방향을 확인했는가", "DXY 추세를 확인했는가", "경상수지·외환보유고를 점검했는가", "환헤지 필요 여부를 결정했는가"],
        "source": "",
    },
    "관세": {
        "lens": [
            ("무역수지", "수출−수입 균형과 품목별 기여도"),
            ("관세율·FTA", "HS코드별 관세율과 FTA 원산지 활용 여부"),
            ("수출경쟁력", "환율·단가·품질 기반 품목 경쟁력"),
            ("통상 환경", "보호무역·통상분쟁·비관세장벽 동향"),
        ],
        "indicators": [
            ("수출증감률", "주력 품목·지역별 증감으로 모멘텀 판단"),
            ("수입증감률", "원자재·중간재 수입은 향후 생산 선행지표"),
            ("무역수지", "흑/적자 추세로 대외건전성 평가"),
            ("관세율", "HS코드·협정세율 적용 여부로 비용 산정"),
        ],
        "risks": ["통상마찰·보복관세", "보호무역 강화", "글로벌 공급망 재편", "환율 변동"],
        "strategy": ["FTA 원산지 활용으로 관세 절감", "HS코드 분류 최적화", "수출시장 다변화"],
        "checklist": ["HS코드·관세율을 확인했는가", "FTA 원산지 기준을 충족하는가", "환율 영향을 반영했는가", "비관세장벽을 점검했는가"],
        "source": "",
    },
    "금융": {
        "lens": [
            ("금리 사이클", "기준금리 방향과 예대마진·대출금리 전이"),
            ("실질금리", "명목금리−물가(CPI)로 자산 매력도 판단"),
            ("가계·신용", "가계부채·연체율·신용 스프레드"),
            ("자산배분", "위험성향별 주식·채권·현금·대안 비중"),
        ],
        "indicators": [
            ("기준금리", "인상기=대출부담↑·예금매력↑, 인하기 반대"),
            ("예금금리", "안전자산 수익률 — 물가 대비 실질수익 확인"),
            ("대출금리", "고정/변동 선택 — 금리 상승기엔 고정 유리"),
            ("CPI", "물가 상승 시 실질구매력·실질금리 하락"),
        ],
        "risks": ["금리 변동", "신용·연체 리스크", "인플레이션", "유동성 경색"],
        "strategy": ["금리 상승기 변동금리 노출 축소", "분산투자·만기 분산", "비상자금 6개월분 확보"],
        "checklist": ["금리 방향을 확인했는가", "대출 고정/변동을 비교했는가", "수수료·세금·중도해지 조건을 확인했는가", "비상자금을 확보했는가"],
        "source": "life-finance-advisor.md",
    },
    "건강": {
        "lens": [
            ("예방·생활습관", "흡연·음주·수면 등 위험요인 관리"),
            ("영양 균형", "탄단지·미량영양소와 일일권장섭취량 대비"),
            ("신체활동", "유산소+근력 운동 강도·빈도"),
            ("정신건강", "스트레스·번아웃·수면의 질"),
            ("정기 검진", "연령·가족력 기반 검진 주기"),
        ],
        "indicators": [
            ("BMI", "18.5~22.9 정상, 23+ 과체중 — 체성분 병행 해석"),
            ("권장섭취량", "활동량·연령별 열량·영양소 기준 대비 점검"),
            ("운동강도", "주 150분 중강도 유산소 + 주 2회 근력 권장"),
        ],
        "risks": ["만성질환(대사·심혈관)", "검증되지 않은 건강정보", "과로·번아웃"],
        "strategy": ["균형 식단·규칙적 운동", "수면 7~8시간 확보", "정기검진·전문의 상담"],
        "checklist": ["BMI가 정상범위인가", "주간 운동량을 점검했는가", "수면·스트레스를 관리하는가", "전문의 상담이 필요한가"],
        "source": "life-health-advisor.md",
    },
    "부동산": {
        "lens": [
            ("가격지수", "매매·전세 가격지수 추이와 전세가율"),
            ("금융·규제", "금리·LTV·DSR 등 대출규제 영향"),
            ("수급", "입주물량·청약경쟁률·미분양"),
            ("정책·세제", "취득·보유·양도세와 규제지역 지정"),
            ("입지", "교통·학군·생활인프라·개발호재"),
        ],
        "indicators": [
            ("매매가격지수", "상승/하락 추세와 지역 차별화"),
            ("전세가율", "전세/매매 비율 — 높으면 갭·실수요 강함"),
            ("청약경쟁률", "수요 강도와 분양시장 온도"),
            ("LTV/DSR", "대출 한도·상환능력 제약"),
        ],
        "risks": ["금리·대출규제", "공급과잉·공실", "정책 변동", "역전세·깡통전세"],
        "strategy": ["실거주 우선·상환능력 내 매수", "입지·학군·교통 중심", "지역·시점 분산"],
        "checklist": ["LTV/DSR 한도를 확인했는가", "전세가율·실거래가를 점검했는가", "입주물량을 확인했는가", "취득·보유·양도세를 계산했는가"],
        "source": "life-realestate-advisor.md",
    },
    "법률": {
        "lens": [
            ("법령·요건", "적용 법령과 권리·의무 성립 요건"),
            ("판례 동향", "유사 사안의 법원 판단 경향"),
            ("시효·기간", "소멸시효·제척기간·제소기간"),
            ("입증·증거", "주장사실의 증거·서류 확보"),
            ("절차·관할", "조정·소송 절차와 관할"),
        ],
        "indicators": [
            ("관련법령", "근거 조문과 요건 충족 여부"),
            ("판례동향", "승소·패소 경향과 쟁점"),
            ("소멸시효", "권리행사 가능 기간(도과 시 권리 소멸)"),
        ],
        "risks": ["소멸시효 도과", "증거 부족", "절차·기간 미준수"],
        "strategy": ["증거 보전·내용증명 우선", "조정·화해 등 비소송 검토", "무료 법률상담(132) 활용"],
        "checklist": ["소멸시효를 확인했는가", "증거·서류를 확보했는가", "관할·절차를 확인했는가", "전문가 상담이 필요한가"],
        "source": "life-legal-advisor.md",
    },
    "교육": {
        "lens": [
            ("목표·진단", "학습 목표와 현재 수준 격차 진단"),
            ("커리큘럼", "목표 역산 로드맵과 단계별 마일스톤"),
            ("학습효율", "인출연습·간격반복 등 메타인지 전략"),
            ("평가·피드백", "형성평가로 약점 보완"),
            ("에듀테크", "온라인·AI 학습도구 활용"),
        ],
        "indicators": [
            ("수강률", "완주·진도율로 학습 지속성 판단"),
            ("합격률", "목표 시험 난이도·준비 적정성"),
            ("학습효율", "투입 시간 대비 성취도"),
        ],
        "risks": ["동기 저하·중도이탈", "정보 과부하", "사교육 과의존"],
        "strategy": ["목표 역산 학습계획", "반복·인출연습 중심", "주기적 형성평가"],
        "checklist": ["목표·기간을 명확히 했는가", "주간 학습량을 점검했는가", "약점 보완 계획이 있는가", "자료 출처를 검증했는가"],
        "source": "life-education-advisor.md",
    },
    "여행": {
        "lens": [
            ("가격·시즌", "항공·숙박 가격과 성수기/비수기"),
            ("환율·물가", "현지 물가·환율로 예산 산정"),
            ("안전", "외교부 여행경보 단계·치안"),
            ("일정 최적화", "동선·이동시간·필수코스 배분"),
        ],
        "indicators": [
            ("항공운임지수", "시즌·노선별 가격 변동"),
            ("숙박가격지수", "지역·등급별 숙박비"),
            ("환율", "현지 통화 환율로 실질 비용 산정"),
        ],
        "risks": ["여행경보·치안", "환율 변동", "항공 지연·결항", "성수기 비용 급등"],
        "strategy": ["비수기·얼리버드 예약", "환전 분산", "여행자보험 가입"],
        "checklist": ["여행경보 단계를 확인했는가", "항공·숙박 가격을 비교했는가", "환율·예산을 점검했는가", "보험·비자를 확인했는가"],
        "source": "life-travel-planner.md",
    },
    "식생활": {
        "lens": [
            ("물가·비용", "외식물가·식재료비 추이"),
            ("영양 균형", "탄단지·식이섬유 균형"),
            ("트렌드", "배달·간편식(HMR)·건강식 수요"),
            ("식품안전", "원산지·유통기한·위생"),
        ],
        "indicators": [
            ("외식물가지수", "외식비 부담 추이"),
            ("식재료가격", "주요 식재료 시세 변동"),
            ("배달주문트렌드", "채널·카테고리 소비 변화"),
        ],
        "risks": ["물가 상승", "식품안전 사고", "영양 불균형"],
        "strategy": ["제철·국산 식재료 활용", "간편식 영양 보완", "배달 빈도 관리"],
        "checklist": ["식재료 가격 추이를 확인했는가", "영양 균형을 점검했는가", "원산지·유통기한을 확인했는가"],
        "source": "life-food-expert.md",
    },
    "쇼핑": {
        "lens": [
            ("소비심리·물가", "소비자심리지수와 물가 영향"),
            ("채널", "온라인/오프라인·라이브커머스 비중"),
            ("프로모션", "할인·적립·카드혜택 주기"),
            ("가성비", "단가·총비용 비교"),
        ],
        "indicators": [
            ("소비자심리지수", "100 기준 — 이상 낙관, 이하 위축"),
            ("온라인거래액", "이커머스 성장·카테고리 트렌드"),
            ("카드소비", "실질 소비 흐름"),
        ],
        "risks": ["충동구매·과소비", "가격 변동", "반품·AS 분쟁"],
        "strategy": ["할인 주기 활용", "최저가·단가 비교", "필요 기반 구매"],
        "checklist": ["최저가를 비교했는가", "할인·적립을 확인했는가", "필요성을 점검했는가", "반품·AS 조건을 확인했는가"],
        "source": "",
    },
    "육아": {
        "lens": [
            ("정책·지원", "보육·돌봄 정책과 수당·바우처"),
            ("발달 단계", "연령별 신체·인지·정서 발달"),
            ("비용", "육아용품·교육·돌봄 비용"),
            ("안전·건강", "예방접종·안전사고 예방"),
        ],
        "indicators": [
            ("합계출산율", "인구·정책 환경 지표"),
            ("보육시설현황", "어린이집·유치원 공급"),
            ("육아용품시장", "용품·서비스 가격 동향"),
        ],
        "risks": ["비용 부담", "정보 과부하", "안전사고"],
        "strategy": ["정부 지원제도 활용", "발달 단계 맞춤 양육", "안전 우선"],
        "checklist": ["지원제도(수당·바우처)를 확인했는가", "발달 단계를 점검했는가", "용품 안전인증을 확인했는가"],
        "source": "",
    },
    "문화": {
        "lens": [
            ("트렌드", "공연·전시·OTT 콘텐츠 동향"),
            ("한류", "K-콘텐츠 수출·글로벌 수요"),
            ("관람 패턴", "장르·연령별 소비"),
            ("비용·접근성", "티켓·예매·지역 인프라"),
        ],
        "indicators": [
            ("공연관람률", "장르별 수요 강도"),
            ("전시참관객", "전시·박람회 흥행"),
            ("한류수출", "콘텐츠 산업 성장"),
        ],
        "risks": ["콘텐츠 편중", "비용 부담", "지역 접근성 격차"],
        "strategy": ["조기예매·할인 활용", "다양한 장르 경험"],
        "checklist": ["공연·전시 일정을 확인했는가", "할인·예매를 확인했는가", "접근성을 점검했는가"],
        "source": "",
    },
    "반려동물": {
        "lens": [
            ("보건·등록", "동물등록·예방접종·건강관리"),
            ("시장·비용", "사료·용품·의료비"),
            ("행동·훈련", "문제행동 예방과 사회화"),
            ("서비스", "펫보험·돌봄·장묘"),
        ],
        "indicators": [
            ("반려동물등록수", "양육 인구·시장 규모"),
            ("펫시장규모", "산업 성장·카테고리"),
            ("사료가격", "양육 비용 동향"),
        ],
        "risks": ["의료비 부담", "유기·분실", "안전사고"],
        "strategy": ["펫보험 가입", "정기 건강검진", "등록·예방접종 준수"],
        "checklist": ["등록·접종을 확인했는가", "사료·영양을 점검했는가", "보험·의료비를 대비했는가"],
        "source": "",
    },
    "화훼": {
        "lens": [
            ("시장·가격", "화훼 도소매 시세·유통"),
            ("트렌드", "플랜테리어·반려식물 수요"),
            ("계절·이벤트", "졸업·기념일 등 수요 집중"),
            ("관리·재배", "광·온도·관수 등 관리 난이도"),
        ],
        "indicators": [
            ("화훼시장규모", "산업 규모·소비 추이"),
            ("플랜테리어 관심도", "홈가드닝 수요 변화"),
        ],
        "risks": ["계절 수요 편중", "관리 난이도", "유통·신선도"],
        "strategy": ["제철 화훼 활용", "관리 쉬운 식물 선택", "이벤트 수요 대응"],
        "checklist": ["제철·시세를 확인했는가", "관리 난이도를 점검했는가", "용도(선물/홈)를 구분했는가"],
        "source": "",
    },
    "유가": {
        "lens": [
            ("유종 스프레드", "WTI·Brent·Dubai 가격차와 정제마진"),
            ("공급", "OPEC+ 감산·증산과 미국 셰일 생산"),
            ("수요·재고", "경기·계절 수요와 원유 재고"),
            ("지정학·환율", "중동 리스크와 원/달러 전이"),
        ],
        "indicators": [
            ("WTI", "미국 서부텍사스유 — 글로벌 벤치마크"),
            ("Brent", "북해유 — 국제 유가 기준"),
            ("Dubai", "중동유 — 국내 도입 비중 높음"),
            ("국내유가", "국제유가 2~3주 시차 전이"),
        ],
        "risks": ["OPEC+ 정책", "지정학 충격", "수요 둔화", "환율 전이"],
        "strategy": ["상승기 비용 헤지", "비용 전가·효율화 검토"],
        "checklist": ["WTI/Brent 추세를 확인했는가", "OPEC+ 일정을 점검했는가", "재고지표를 확인했는가", "환율 전이를 반영했는가"],
        "source": "",
    },
    "운송": {
        "lens": [
            ("운임 지수", "BDI(건화물)·SCFI(컨테이너) 운임"),
            ("유가 연동", "벙커유·유가가 운임 원가에 전이"),
            ("물동량", "교역량·항만 처리량"),
            ("공급망", "선복·항만 적체·공급망 차질"),
        ],
        "indicators": [
            ("BDI", "건화물 운임 — 글로벌 교역 선행지표"),
            ("컨테이너운임", "공급망·소비재 물류 비용"),
            ("항공화물량", "고부가·긴급 화물 수요"),
        ],
        "risks": ["운임 변동", "유가 상승", "공급망 차질", "항만 적체"],
        "strategy": ["장·단기 운임계약 분산", "유가 헤지", "공급망 다변화"],
        "checklist": ["BDI/SCFI 추세를 확인했는가", "유가 영향을 점검했는가", "공급망 리스크를 점검했는가"],
        "source": "",
    },
    "분쟁": {
        "lens": [
            ("지정학 리스크", "분쟁 강도·확전 가능성"),
            ("자원 충격", "에너지·곡물 등 원자재 공급 영향"),
            ("안보·동맹", "동맹·제재·외교 동향"),
            ("시장 파급", "안전자산 선호·변동성 전이"),
        ],
        "indicators": [
            ("지정학리스크지수", "리스크 고조/완화 추세"),
            ("분쟁영향권국가", "공급망·교역 노출도"),
            ("난민현황", "인도적·사회적 파급"),
        ],
        "risks": ["에너지·원자재 충격", "공급망 차질", "안전·여행 위험", "시장 변동성"],
        "strategy": ["안전자산 비중 조정", "공급망 다변화", "여행·교역 리스크 관리"],
        "checklist": ["분쟁 지역·영향권을 확인했는가", "에너지/곡물 가격 영향을 점검했는가", "안전·여행경보를 확인했는가"],
        "source": "",
    },
    "사업": {
        "lens": [
            ("창업 동향", "창업·폐업 추이와 업종 트렌드"),
            ("자금조달", "VC·대출·정부지원 환경"),
            ("시장·경쟁", "수요·경쟁강도·진입장벽"),
            ("규제·정책", "인허가·규제·세제"),
            ("수익성", "비즈니스모델·손익분기·현금흐름"),
        ],
        "indicators": [
            ("창업건수", "창업 활력·업종 쏠림"),
            ("폐업률", "시장 경쟁·생존율"),
            ("VC투자규모", "자금조달 환경·투심"),
        ],
        "risks": ["자금경색", "경쟁 심화", "규제 변동", "수요 변동"],
        "strategy": ["린(Lean) 시장 검증", "정부지원·R&D 활용", "현금흐름 중심 운영"],
        "checklist": ["시장·경쟁을 분석했는가", "자금조달 계획이 있는가", "규제·인허가를 확인했는가", "손익분기를 점검했는가"],
        "source": "",
    },
    "default": {
        "lens": [
            ("데이터 신뢰도", "출처·기준일·표본의 대표성 검증"),
            ("시계열 추세", "방향성·기울기와 추세 지속성"),
            ("변동성·이상치", "표준편차·CV와 급변 구간 식별"),
            ("뉴스 감성", "긍정/부정 논조와 키워드 흐름"),
            ("선행지표", "상관·선행 관계로 방향 예측"),
        ],
        "indicators": [
            ("관심도지수", "검색·언급 빈도로 수요 강도 판단"),
            ("뉴스빈도", "이슈화 정도와 모멘텀"),
            ("검색트렌드", "대중 관심의 선행 신호"),
        ],
        "risks": ["데이터 편향", "단기 노이즈", "외부 충격"],
        "strategy": ["다출처 교차검증", "추세·변동성 병행 분석", "정기 모니터링"],
        "checklist": ["출처·기준일을 확인했는가", "추세 방향을 점검했는가", "이상치 원인을 분석했는가", "추가 데이터를 확보했는가"],
        "source": "life-data-analyst.md",
    },
}


def _domain_dd(domain) -> dict:
    """매칭된 도메인 프로필에서 심층 분석 지식 묶음을 반환."""
    key = domain.get("_key", "default") if isinstance(domain, dict) else "default"
    return DOMAIN_DEEPDIVE.get(key) or DOMAIN_DEEPDIVE["default"]


def _domain_narrative(domain, query, terms, figs, stats) -> str:
    """검색어·추출수치·통계를 분야 관점에 연결한 맞춤 심층 해석 서술."""
    dd = _domain_dd(domain)
    framework = domain.get("framework", "데이터 기반 분석") if isinstance(domain, dict) else "데이터 기반 분석"
    term_str = ", ".join(terms[:5]) if terms else (query or "검색 주제")
    parts = [f"검색 주제 「{query}」의 핵심어({term_str})를 {framework} 관점에서 종합하면 다음과 같습니다."]
    lens = dd.get("lens", [])
    if lens:
        s = f"우선 '{lens[0][0]}' 측면에서 {lens[0][1]}."
        if len(lens) > 1:
            s += f" 또한 '{lens[1][0]}' 측면에서 {lens[1][1]}."
        parts.append(s)
    if figs:
        parts.append(
            f"자료에서 추출된 수치({' · '.join(figs[:5])})는 위 관점의 정량 근거이며, "
            f"단일 시점이 아닌 추세·맥락으로 해석해야 합니다."
        )
    if stats:
        risk0 = dd.get("risks", ["주요 리스크"])[0]
        parts.append(
            f"수집 트렌드는 {stats['trend_dir']}({stats['pct_change']:+.1f}%)·변동성 {stats['volatility']}"
            f"(CV {stats['cv']:.1f}%) 수준으로, '{risk0}' 점검이 우선됩니다."
        )
    return " ".join(parts)


def _render_domain_deepdive_word(doc, domain, query, stats, figs, terms, Pt, RGBColor):
    """[Word] 검색 종류 맞춤 분야 전문 심층 분석 블록 — lens·지표·연계해석·리스크·체크리스트."""
    from docx.enum.table import WD_TABLE_ALIGNMENT
    dd = _domain_dd(domain)
    icon = domain.get("icon", "📊") if isinstance(domain, dict) else "📊"
    framework = domain.get("framework", "데이터 기반 분석") if isinstance(domain, dict) else "데이터 기반 분석"

    p = doc.add_paragraph()
    r = p.add_run(f"▶ {icon} 분야 전문가 심층 분석 — {framework}")
    r.bold = True
    r.font.size = Pt(11)

    # 1) 전문가 분석 관점
    p = doc.add_paragraph()
    r = p.add_run("· 전문가 분석 관점")
    r.bold = True
    r.font.size = Pt(10)
    for title, desc in dd.get("lens", []):
        bp = doc.add_paragraph(style="List Bullet")
        rt = bp.add_run(f"{title} — ")
        rt.bold = True
        rt.font.size = Pt(9)
        rd = bp.add_run(desc)
        rd.font.size = Pt(9)

    # 2) 핵심 지표 해석 가이드
    inds = dd.get("indicators", [])
    if inds:
        p = doc.add_paragraph()
        r = p.add_run("· 핵심 지표 해석 가이드")
        r.bold = True
        r.font.size = Pt(10)
        tbl = doc.add_table(rows=1 + len(inds), cols=2, style="Light Shading Accent 1")
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        tbl.rows[0].cells[0].text = "지표"
        tbl.rows[0].cells[1].text = "해석 기준"
        for c in tbl.rows[0].cells:
            for pp in c.paragraphs:
                for rr in pp.runs:
                    rr.bold = True
        for i, (ind, interp) in enumerate(inds):
            tbl.rows[i + 1].cells[0].text = ind
            tbl.rows[i + 1].cells[1].text = interp

    # 3) 검색어 연계 심층 해석
    p = doc.add_paragraph()
    r = p.add_run("· 검색어 연계 심층 해석")
    r.bold = True
    r.font.size = Pt(10)
    doc.add_paragraph(_domain_narrative(domain, query, terms, figs, stats))

    # 4) 분야별 리스크 점검
    risks = dd.get("risks", [])
    if risks:
        p = doc.add_paragraph()
        r = p.add_run("· 분야별 리스크 점검")
        r.bold = True
        r.font.size = Pt(10)
        doc.add_paragraph("   " + "  ·  ".join(risks))

    # 5) 실행 체크리스트
    checklist = dd.get("checklist", [])
    if checklist:
        p = doc.add_paragraph()
        r = p.add_run("· 실행 체크리스트")
        r.bold = True
        r.font.size = Pt(10)
        for c in checklist:
            doc.add_paragraph(f"   ☐ {c}")

    # 근거 출처
    src = dd.get("source")
    if src:
        sp = doc.add_paragraph(f"   ※ 분석 프레임워크 근거: 02_생활정보/00_Skill_MD/{src}")
        for rn in sp.runs:
            rn.font.size = Pt(7)
            rn.font.color.rgb = RGBColor(0x99, 0x99, 0x99)


def _render_domain_strategy_word(doc, domain, Pt, RGBColor):
    """[Word] 전문가 종합 인사이트에 추가하는 분야별 전략 관점 + 모니터링 지표."""
    dd = _domain_dd(domain)
    strat = dd.get("strategy", [])
    inds = dd.get("indicators", [])
    if not strat and not inds:
        return
    framework = domain.get("framework", "분야") if isinstance(domain, dict) else "분야"
    doc.add_paragraph("")
    p = doc.add_paragraph()
    r = p.add_run(f"▶ 분야별 전략 관점 — {framework}")
    r.bold = True
    r.font.size = Pt(11)
    for s in strat:
        doc.add_paragraph(f"   • {s}")
    if inds:
        names = ", ".join(ind for ind, _ in inds)
        mp = doc.add_paragraph(f"   핵심 모니터링 지표: {names}")
        for rn in mp.runs:
            rn.font.size = Pt(9)
            rn.font.color.rgb = RGBColor(0x55, 0x55, 0x55)


def _domain_deepdive_text_lines(domain, query, stats, figs, terms) -> list:
    """[Text] 분야 전문 심층 분석을 텍스트 라인 목록으로 반환."""
    dd = _domain_dd(domain)
    framework = domain.get("framework", "데이터 기반 분석") if isinstance(domain, dict) else "데이터 기반 분석"
    lines = [f"  [ 분야 전문가 심층 분석 — {framework} ]"]
    for title, desc in dd.get("lens", []):
        lines.append(f"    · {title}: {desc}")
    inds = dd.get("indicators", [])
    if inds:
        lines.append("    핵심 지표 해석:")
        for ind, interp in inds:
            lines.append(f"      - {ind}: {interp}")
    lines.append("    연계 해석: " + _domain_narrative(domain, query, terms, figs, stats))
    risks = dd.get("risks", [])
    if risks:
        lines.append("    리스크: " + " · ".join(risks))
    strat = dd.get("strategy", [])
    if strat:
        lines.append("    전략: " + " / ".join(strat))
    checklist = dd.get("checklist", [])
    if checklist:
        lines.append("    체크리스트: " + " / ".join(checklist))
    src = dd.get("source")
    if src:
        lines.append(f"    ※ 근거: 00_Skill_MD/{src}")
    return lines


def _is_text_similar(text1: str, text2: str, threshold: float = 0.6) -> bool:
    """Check if two texts are too similar by word overlap ratio."""
    if not text1 or not text2:
        return False
    # Normalize: remove punctuation, lowercase
    import re as _re
    norm1 = _re.sub(r"[^\w\s]", "", text1).split()
    norm2 = _re.sub(r"[^\w\s]", "", text2).split()
    if not norm1 or not norm2:
        return False
    set1, set2 = set(norm1), set(norm2)
    overlap = len(set1 & set2)
    return overlap / min(len(set1), len(set2)) >= threshold


# ═══════════════════════════════════════════════════════════════════════════════
# Statistical helpers
# ═══════════════════════════════════════════════════════════════════════════════

def _calc_statistics(values: list) -> dict:
    """Calculate comprehensive statistics for a trend series."""
    if not values or len(values) < 2:
        return {}
    n = len(values)
    mean_val = sum(values) / n
    variance = sum((v - mean_val) ** 2 for v in values) / (n - 1)
    std_dev = math.sqrt(variance) if variance > 0 else 0
    max_val = max(values)
    min_val = min(values)
    range_val = max_val - min_val

    # Simple moving average (3-period)
    sma3 = []
    for i in range(2, n):
        sma3.append(sum(values[i - 2:i + 1]) / 3)

    # Trend direction via linear regression slope
    x_mean = (n - 1) / 2
    numerator = sum((i - x_mean) * (v - mean_val) for i, v in enumerate(values))
    denominator = sum((i - x_mean) ** 2 for i in range(n))
    slope = numerator / denominator if denominator != 0 else 0

    # Coefficient of variation (volatility proxy)
    cv = (std_dev / mean_val * 100) if mean_val != 0 else 0

    # Change metrics
    first_val = values[0]
    last_val = values[-1]
    abs_change = last_val - first_val
    pct_change = ((last_val - first_val) / first_val * 100) if first_val != 0 else 0

    return {
        "mean": mean_val,
        "std_dev": std_dev,
        "max": max_val,
        "min": min_val,
        "range": range_val,
        "cv": cv,
        "slope": slope,
        "sma3": sma3,
        "abs_change": abs_change,
        "pct_change": pct_change,
        "first": first_val,
        "last": last_val,
        "trend_dir": "상승" if slope > 0 else ("하락" if slope < 0 else "보합"),
        "volatility": "높음" if cv > 15 else ("보통" if cv > 5 else "낮음"),
    }


def _extract_news_themes(news: list) -> list[str]:
    """Extract common themes/keywords from news titles for flow analysis."""
    if not news:
        return []
    # Simple keyword frequency
    from collections import Counter
    stop_words = {"및", "등", "위한", "대한", "관련", "통해", "따라", "에서", "으로", "이번", "최근", "올해"}
    words = []
    for n in news:
        title = _clean_text(n.get("title", ""))
        for w in title.split():
            w = w.strip(".,!?()[]\"'·…")
            if len(w) >= 2 and w not in stop_words:
                words.append(w)
    counter = Counter(words)
    return [w for w, _ in counter.most_common(8)]


def _news_flow_summary(news: list, query: str) -> str:
    """Generate a narrative flow analysis from news articles."""
    if not news:
        return ""
    themes = _extract_news_themes(news)
    n_count = len(news)

    sources = set()
    for n in news:
        s = n.get("source", "")
        if s:
            sources.add(s)

    flow = f"수집된 {n_count}건의 뉴스를 종합 분석한 결과, "
    if themes:
        flow += f"'{query}' 관련 핵심 키워드는 【{'、'.join(themes[:5])}】 등으로 나타났습니다. "
    if sources:
        flow += f"주요 보도 매체는 {', '.join(list(sources)[:4])} 등입니다. "

    # Simple sentiment heuristic from title keywords
    pos_words = ["상승", "호조", "성장", "확대", "개선", "증가", "강세", "반등", "돌파", "긍정"]
    neg_words = ["하락", "감소", "위기", "약세", "우려", "축소", "급락", "둔화", "침체", "적자"]
    pos_count = 0
    neg_count = 0
    for n in news:
        title = n.get("title", "")
        for pw in pos_words:
            if pw in title:
                pos_count += 1
        for nw in neg_words:
            if nw in title:
                neg_count += 1

    total_sentiment = pos_count + neg_count
    if total_sentiment > 0:
        pos_ratio = pos_count / total_sentiment * 100
        neg_ratio = neg_count / total_sentiment * 100
        if pos_ratio > 60:
            flow += f"전반적 논조는 긍정적(긍정 {pos_ratio:.0f}% vs 부정 {neg_ratio:.0f}%)이며, 시장/분야 전망에 대한 낙관론이 우세합니다."
        elif neg_ratio > 60:
            flow += f"전반적 논조는 부정적(부정 {neg_ratio:.0f}% vs 긍정 {pos_ratio:.0f}%)이며, 리스크 요인에 대한 경계감이 감지됩니다."
        else:
            flow += f"긍정(긍정 {pos_ratio:.0f}%)과 부정(부정 {neg_ratio:.0f}%) 논조가 혼재하며, 불확실성이 높은 상황입니다."
    else:
        flow += "뉴스 논조는 중립적이며, 주요 방향성 전환 신호는 아직 감지되지 않았습니다."

    return flow


# ═══════════════════════════════════════════════════════════════════════════════
# Query-focused analysis (검색 주제 맞춤 상세 코너)
# ═══════════════════════════════════════════════════════════════════════════════

_QUERY_STOPWORDS = {
    "분석", "세부", "기반", "대한", "관련", "최근", "동향", "현황", "정보", "추세",
    "리포트", "보고서", "내용", "대해", "위한", "중심", "전반", "전체", "상세", "주요",
    "어떻게", "무엇", "그리고", "또한", "전망", "예측", "방안", "비교", "검토",
}


def _query_key_terms(query: str) -> list[str]:
    """질문에서 핵심 검색어(2자+ 한글/영문/숫자)를 추출 — 불용어 제외, 순서 보존."""
    import re as _re
    toks = _re.findall(r"[가-힣A-Za-z0-9]{2,}", query or "")
    terms, seen = [], set()
    for t in toks:
        tl = t.lower()
        if tl in _QUERY_STOPWORDS or tl in seen:
            continue
        seen.add(tl)
        terms.append(t)
    return terms[:10]


def _detect_time_scope(query: str) -> str:
    """질문의 기간 의도 감지 → 분석 범위 라벨('1년'·'6개월' 등). 없으면 ''."""
    import re as _re
    q = query or ""
    m = _re.search(r"(\d+)\s*(년|개월|분기|주|일)", q)
    if m:
        return f"{m.group(1)}{m.group(2)}"
    for k, label in [("연간", "1년"), ("1년", "1년"), ("장기", "장기"),
                     ("분기", "분기"), ("월간", "1개월"), ("주간", "1주")]:
        if k in q:
            return label
    return ""


def _score_item_relevance(item: dict, terms: list[str]) -> int:
    """뉴스/웹 항목 텍스트와 검색어 일치 빈도 점수."""
    if not terms:
        return 0
    text = (str(item.get("title", "")) + " " +
            str(item.get("snippet", "") or item.get("source", ""))).lower()
    return sum(text.count(t.lower()) for t in terms)


def _extract_figures_from_items(items: list, limit: int = 8) -> list[str]:
    """뉴스/웹 본문에서 핵심 수치(%, 통화, 배수·포인트, 연도)를 자동 추출."""
    import re as _re
    pats = [
        r"[+\-]?\d[\d,\.]*\s?%",
        r"\d[\d,\.]*\s?(?:조|억|만)?\s?(?:원|달러|USD|\$|엔|위안)",
        r"\d[\d,\.]*\s?(?:배|bp|포인트|건|명|개월|개)",
        r"20\d{2}\s?년",
    ]
    out, seen = [], set()
    for it in items:
        text = str(it.get("title", "")) + " " + str(it.get("snippet", "") or "")
        for p in pats:
            for m in _re.findall(p, text):
                s = m.strip()
                key = _re.sub(r"\s+", "", s)
                if s and key not in seen and len(s) <= 24:
                    seen.add(key)
                    out.append(s)
                    if len(out) >= limit:
                        return out
    return out


def _rank_relevant_items(query: str, news: list, web: list, top: int = 6):
    """검색어 관련성으로 뉴스+웹을 순위화. (relevant_items, key_terms, n_matched, time_scope)."""
    terms = _query_key_terms(query)
    scope = _detect_time_scope(query)
    items = []
    for n in (news or []):
        items.append({"_kind": "뉴스", **n, "_score": _score_item_relevance(n, terms)})
    for w in (web or []):
        items.append({"_kind": "웹", **w, "_score": _score_item_relevance(w, terms)})
    items.sort(key=lambda x: x["_score"], reverse=True)
    n_matched = len([it for it in items if it["_score"] > 0])
    relevant = [it for it in items if it["_score"] > 0][:top] or items[:max(3, top // 2)]
    return relevant, terms, n_matched, scope


def _quick_sentiment(news) -> str:
    """뉴스 제목 기반 빠른 감성 판정(긍정/부정/중립)."""
    if not news:
        return "중립"
    pos_words = ["상승", "호조", "성장", "확대", "개선", "증가", "강세", "반등", "돌파", "긍정", "회복", "흑자", "수혜"]
    neg_words = ["하락", "감소", "위기", "약세", "우려", "축소", "급락", "둔화", "침체", "적자", "붕괴", "폭락", "리스크"]
    pc = nc = 0
    for n in news:
        t = n.get("title", "")
        pc += sum(1 for w in pos_words if w in t)
        nc += sum(1 for w in neg_words if w in t)
    if pc > nc * 1.2:
        return "긍정"
    if nc > pc * 1.2:
        return "부정"
    return "중립"


def _topic_synthesis_paras(query, news, web, relevant, terms, figs, scope, n_matched, domain, stats) -> list:
    """검색 주제에 대한 상세 내용을 수집자료로 합성 — 개요·논점·논조·수치·기간·평가 다단락.
    LLM 없이 수집된 뉴스/웹/수치/통계를 재구성한 근거 기반 서술(원문 검증 권장)."""
    dd = _domain_dd(domain)
    framework = domain.get("framework", "데이터 기반 분석") if isinstance(domain, dict) else "데이터 기반 분석"
    n_news = len(news or [])
    n_web = len(web or [])
    themes = _extract_news_themes(news) if news else []
    lens = dd.get("lens", [])
    risks = dd.get("risks", [])
    paras = []

    # 1) 개요 + 핵심 논점
    p1 = f"수집된 뉴스 {n_news}건·웹 {n_web}건(주제 직접 매칭 {n_matched}건)을 '{query}' 관점에서 종합했습니다."
    if themes:
        p1 += f" 자료에서 반복적으로 등장한 핵심 논점은 【{'、'.join(themes[:6])}】 중심입니다."
    paras.append(p1)

    # 2) 논조/감성 흐름 (뉴스 흐름 요약 재사용)
    flow = _news_flow_summary(news, query) if news else ""
    if flow:
        paras.append(flow)

    # 3) 핵심 수치 연계
    if figs:
        l0 = lens[0][0] if lens else "핵심 지표"
        l1 = lens[1][0] if len(lens) > 1 else l0
        paras.append(
            f"자료에서 확인된 핵심 수치는 {' · '.join(figs[:6])}입니다. "
            f"이는 {framework} 관점에서 '{l0}'·'{l1}' 판단의 정량 근거이며, "
            f"단일 수치가 아닌 추세·맥락으로 해석해야 합니다."
        )

    # 4) 기간(scope) 기반 추세 해석
    if scope:
        if stats:
            paras.append(
                f"질문하신 '{scope}' 기간 관점에서, 수집 트렌드는 {stats['trend_dir']}"
                f"({stats['pct_change']:+.1f}%)·변동성 {stats['volatility']}(CV {stats['cv']:.1f}%)로 나타났습니다. "
                f"기간 전반의 최고·최저 시점과 변동 요인을 시계열로 연결해 방향성을 판단해야 합니다."
            )
        else:
            paras.append(
                f"질문하신 '{scope}' 기간 분석에는 해당 기간의 시계열 데이터가 필요합니다. "
                f"현재 트렌드 표본이 부족하므로 차트 기간을 '{scope}'에 맞춰 확장해 재분석하면 정확도가 높아집니다."
            )

    # 5) 주제 평가 (종합)
    sent = _quick_sentiment(news)
    sent_label = {"긍정": "긍정적 신호가 우세", "부정": "부정적·경계 신호가 우세", "중립": "방향성이 혼재"}.get(sent, "방향성이 혼재")
    p5 = f"이상을 종합하면, '{query}'에 대해 현재 공개 자료는 {sent_label}합니다. "
    if lens:
        watch = ", ".join(t for t, _ in lens[:3])
        p5 += f"추가로 {watch} 지표를 점검하고, "
    if risks:
        r1 = risks[1] if len(risks) > 1 else risks[0]
        p5 += f"특히 {risks[0]}·{r1}에 유의해 시나리오별로 대응할 것을 권고합니다."
    paras.append(p5)

    return paras


def _add_query_focus_section_word(doc, query, news, web, Pt, RGBColor, domain=None, stats=None):
    """[Word] 검색 주제 맞춤 상세 코너 — 질문 의도에 맞춰 관련 자료를 선별·종합 + 분야 전문 심층 분석."""
    relevant, terms, n_matched, scope = _rank_relevant_items(query, news, web)

    doc.add_paragraph(
        f"검색하신 주제 「{query}」에 초점을 맞춰, 수집된 뉴스·웹 자료 중 관련성이 높은 항목을 "
        f"선별·종합하고 핵심 수치를 추출했습니다. (기본 분석 포맷은 이하 섹션에서 그대로 제공됩니다.)"
    )
    if terms:
        doc.add_paragraph("▶ 핵심 검색어: " + "  ".join(f"【{t}】" for t in terms[:6]))
    if scope:
        doc.add_paragraph(
            f"▶ 질문하신 기간 의도: '최근 {scope}' — 아래 자료의 시점·수치를 해당 기간 관점에서 종합 해석합니다."
        )

    figs = _extract_figures_from_items(relevant, limit=8)

    # ▶ 주제 핵심 논점 종합 (검색 주제 상세 내용 — 수집자료 합성)
    p = doc.add_paragraph()
    r = p.add_run("▶ 주제 핵심 논점 종합")
    r.bold = True
    r.font.size = Pt(11)
    for para in _topic_synthesis_paras(query, news, web, relevant, terms, figs, scope, n_matched, domain, stats):
        pp = doc.add_paragraph(para)
        pp.paragraph_format.space_after = Pt(4)

    # 주제 관련 핵심 자료
    p = doc.add_paragraph()
    r = p.add_run("▶ 주제 관련 핵심 자료 (관련도 순)")
    r.bold = True
    r.font.size = Pt(11)
    if relevant:
        for it in relevant:
            title = _clean_text(it.get("title", ""))[:140]
            link = it.get("link", "")
            snip = _clean_text(it.get("snippet", "") or "")[:170]
            pp = doc.add_paragraph(style="List Bullet")
            tag = pp.add_run(f"[{it['_kind']}] ")
            tag.bold = True
            tag.font.size = Pt(9)
            if link:
                _add_hyperlink(pp, title, link)
            else:
                pp.add_run(title).font.size = Pt(9)
            if snip:
                sp = doc.add_paragraph(f"      {snip}")
                for rn in sp.runs:
                    rn.font.size = Pt(8)
                    rn.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    else:
        doc.add_paragraph(
            "주제와 직접 일치하는 자료가 충분치 않습니다. 검색어를 더 구체화하면(예: 종목·기간·지표 명시) "
            "주제 맞춤 분석 정확도가 높아집니다."
        )

    # 추출 수치 (위에서 계산한 figs 재사용)
    if figs:
        p = doc.add_paragraph()
        r = p.add_run("▶ 자료에서 추출한 핵심 수치·사실")
        r.bold = True
        r.font.size = Pt(11)
        doc.add_paragraph("   " + "  ·  ".join(figs))
        note = doc.add_paragraph(
            "   ※ 위 수치는 기사·검색 결과 본문에서 자동 추출한 참고값입니다. 정확성은 각 원문 링크로 검증하십시오."
        )
        for rn in note.runs:
            rn.font.size = Pt(8)
            rn.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    # 주제 맞춤 시사점
    p = doc.add_paragraph()
    r = p.add_run("▶ 주제 맞춤 시사점")
    r.bold = True
    r.font.size = Pt(11)
    insight = f"'{query}' 키워드에 직접 매칭된 자료는 {n_matched}건입니다. "
    if scope:
        insight += (
            f"질문하신 '{scope}' 관점에서는 단일 시점이 아닌 기간 전반의 방향성과 변동 요인을 "
            f"함께 검토해야 하며, 위 자료의 시점·수치를 시계열로 연결해 추세를 판단하시기 바랍니다. "
        )
    if n_matched == 0:
        insight += "현재 검색어와 직접 매칭이 약하므로, 더 구체적인 키워드로 재검색을 권장합니다. "
    insight += "정량 통계·감성·트렌드 등 표준 분석은 이하 섹션을 참조하세요."
    doc.add_paragraph(insight)

    # ── 분야 전문 심층 분석 (검색 종류 맞춤) ──
    if domain is not None:
        doc.add_paragraph("")
        _render_domain_deepdive_word(doc, domain, query, stats, figs, terms, Pt, RGBColor)


# ═══════════════════════════════════════════════════════════════════════════════
# Standard Word TOC field (표준 목차 — 헤딩 기반 자동 페이지번호)
# ═══════════════════════════════════════════════════════════════════════════════

def _set_update_fields_on_open(doc):
    """Word가 문서를 열 때 TOC 등 필드를 자동 갱신하도록 설정."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    try:
        settings = doc.settings.element
        if settings.find(qn("w:updateFields")) is None:
            uf = OxmlElement("w:updateFields")
            uf.set(qn("w:val"), "true")
            settings.append(uf)
    except Exception:
        pass


def _add_toc_field(doc, levels: str = "2-3"):
    """표준 Word TOC 필드 삽입 — 헤딩 스타일에서 항목·페이지번호 자동 생성."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    para = doc.add_paragraph()
    run = para.add_run()
    r = run._r
    fb = OxmlElement("w:fldChar")
    fb.set(qn("w:fldCharType"), "begin")
    fb.set(qn("w:dirty"), "true")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f' TOC \\o "{levels}" \\h \\z \\u '
    fs = OxmlElement("w:fldChar")
    fs.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "[목차] 항목 위에서 우클릭 → '필드 업데이트(Update Field)'를 선택하면 페이지 번호가 채워집니다."
    fe = OxmlElement("w:fldChar")
    fe.set(qn("w:fldCharType"), "end")
    for el in (fb, instr, fs, placeholder, fe):
        r.append(el)
    return para


def _add_pageref_run(paragraph, bookmark, Pt, RGBColor):
    """PAGEREF 필드 run 추가 — Word가 열릴 때 북마크의 실제 페이지번호로 갱신(하이퍼링크 아님)."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    run = paragraph.add_run()
    run.font.size = Pt(10.5)
    r = run._r
    fb = OxmlElement("w:fldChar")
    fb.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f' PAGEREF {bookmark} '  # \\h 미사용 = 페이지번호만(링크 아님)
    fs = OxmlElement("w:fldChar")
    fs.set(qn("w:fldCharType"), "separate")
    t = OxmlElement("w:t")
    t.text = "·"  # Word 갱신 전 placeholder (갱신 후 페이지번호)
    fe = OxmlElement("w:fldChar")
    fe.set(qn("w:fldCharType"), "end")
    for el in (fb, instr, fs, t, fe):
        r.append(el)


def _add_standard_toc(doc, entries, Pt, RGBColor):
    """표준 Word TOC 필드 + 캐시된 일반 텍스트 항목(번호·제목·점선리더·PAGEREF 페이지번호).

    entries: list of (label, title, bookmark). **하이퍼링크 없는 표준 목차** — 항목은 검은
    일반 텍스트, 우측에 점선 리더 + 페이지번호. Word에서 열면 헤딩 기반으로 목차·페이지번호가
    자동 갱신된다(\\h 미사용이라 항목이 파란 링크로 바뀌지 않음). 하이퍼링크는 출처/원문에만 사용.
    """
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    # ── TOC 필드 시작 (begin + instrText + separate) — \h(하이퍼링크) 미사용 = 표준 목차 ──
    p = doc.add_paragraph()
    r = p.add_run()._r
    fb = OxmlElement("w:fldChar")
    fb.set(qn("w:fldCharType"), "begin")
    fb.set(qn("w:dirty"), "true")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "2-2" \\z \\u '
    fs = OxmlElement("w:fldChar")
    fs.set(qn("w:fldCharType"), "separate")
    for el in (fb, instr, fs):
        r.append(el)

    # ── 캐시된 목차 항목 (필드 갱신 전까지 표시) — 일반 텍스트(하이퍼링크 아님) ──
    for label, title, bm in entries:
        ep = doc.add_paragraph()
        pPr = ep._p.get_or_add_pPr()
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "right")
        tab.set(qn("w:leader"), "dot")
        tab.set(qn("w:pos"), "9072")  # ~16cm 우측 정렬 점선 리더
        tabs.append(tab)
        pPr.append(tabs)
        run = ep.add_run(f"{label}. {title}")
        run.font.size = Pt(10.5)
        run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
        tr = OxmlElement("w:r")
        tr.append(OxmlElement("w:tab"))
        ep._p.append(tr)
        _add_pageref_run(ep, bm, Pt, RGBColor)

    # ── TOC 필드 종료 ──
    endp = doc.add_paragraph()
    er = endp.add_run()._r
    fe = OxmlElement("w:fldChar")
    fe.set(qn("w:fldCharType"), "end")
    er.append(fe)


def _plan_individual_sections(query, news, web, trend, table_data, youtube):
    """개별 리포트의 섹션 순서·번호·북마크 플랜 — TOC와 본문 헤딩 번호를 일치시킨다.
    조건은 _gen_word 본문의 섹션 생성 조건과 정확히 동일해야 한다."""
    plan = [("분석 개요 (Executive Summary)", "sec_summary")]
    if query:
        plan.append(("검색 주제 심층 분석", "sec_focus"))
    if trend:
        plan.append(("트렌드 분석 및 통계", "sec_trend"))
    if trend and len(trend) >= 3:
        plan.append(("변화율 분석 및 이상치 탐지", "sec_change_rate"))
    if table_data:
        plan.append(("데이터 현황 및 통계 분석", "sec_table_data"))
    if news:
        plan.append(("뉴스 심층 분석 (감성·키워드·출처)", "sec_news_analysis"))
        plan.append(("주요 뉴스 목록", "sec_news_list"))
    if web:
        plan.append(("웹 검색 결과 분석", "sec_web"))
    if youtube:
        plan.append(("관련 YouTube 영상", "sec_youtube"))
    plan.append(("전문가 종합 인사이트", "sec_expert"))
    plan.append(("참고문헌 (References)", "sec_refs"))
    return [(i + 1, t, b) for i, (t, b) in enumerate(plan)]


# ═══════════════════════════════════════════════════════════════════════════════
# Chart generation
# ═══════════════════════════════════════════════════════════════════════════════

def _make_trend_chart(trend, query, stats=None) -> io.BytesIO:
    """Generate an enhanced trend chart with statistics overlay."""
    plt = _setup_korean_font()

    dates = [str(r.get("Date", "")) for r in trend]
    values = [r.get("Trend", 0) for r in trend]

    fig, ax = plt.subplots(figsize=(7.5, 4))

    # Main line + fill
    ax.fill_between(range(len(dates)), values, alpha=0.15, color="#1976D2")
    ax.plot(range(len(dates)), values, marker="o", color="#1976D2",
            linewidth=2.5, markersize=7, zorder=5, label="트렌드 지수")

    # Data labels
    for i, v in enumerate(values):
        ax.annotate(f"{v:,.0f}", (i, v), textcoords="offset points",
                    xytext=(0, 12), ha="center", fontsize=8, fontweight="bold",
                    color="#1976D2")

    # Mean line
    if stats and stats.get("mean"):
        ax.axhline(y=stats["mean"], color="#E65100", linestyle="--",
                   linewidth=1, alpha=0.7, label=f'평균: {stats["mean"]:,.1f}')

    # SMA-3 overlay
    if stats and stats.get("sma3") and len(stats["sma3"]) >= 2:
        sma_x = list(range(2, 2 + len(stats["sma3"])))
        ax.plot(sma_x, stats["sma3"], color="#4CAF50", linewidth=1.5,
                linestyle="-.", alpha=0.8, label="3일 이동평균")

    # Min/Max markers
    if len(values) >= 2:
        max_idx = values.index(max(values))
        min_idx = values.index(min(values))
        ax.annotate("▲ HIGH", (max_idx, values[max_idx]),
                    textcoords="offset points", xytext=(0, 22),
                    ha="center", fontsize=7, color="#D32F2F", fontweight="bold")
        ax.annotate("▼ LOW", (min_idx, values[min_idx]),
                    textcoords="offset points", xytext=(0, -18),
                    ha="center", fontsize=7, color="#1565C0", fontweight="bold")

    ax.set_xticks(range(len(dates)))
    ax.set_xticklabels(dates, rotation=30, ha="right", fontsize=9)
    ax.set_ylabel("트렌드 지수", fontsize=10)
    ax.set_title(f"'{query}' 최근 트렌드 추이 분석", fontsize=12, fontweight="bold")
    ax.grid(axis="y", alpha=0.2)
    ax.grid(axis="x", alpha=0.1)
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    ax.legend(loc="upper right", fontsize=8, framealpha=0.8)
    fig.tight_layout()

    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=150, bbox_inches="tight")
    plt.close(fig)
    buf.seek(0)
    return buf


def _setup_korean_font():
    """Configure matplotlib to use Korean-capable font."""
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    import matplotlib.font_manager as fm
    available_fonts = {f.name for f in fm.fontManager.ttflist}
    for fname in ["Malgun Gothic", "NanumGothic", "Noto Sans KR", "Noto Sans CJK KR",
                   "Noto Sans CJK JP", "AppleGothic", "DejaVu Sans"]:
        if fname in available_fonts:
            plt.rcParams["font.family"] = fname
            break
    plt.rcParams["axes.unicode_minus"] = False
    return plt


def _make_keyword_freq_chart(news: list, query: str) -> io.BytesIO | None:
    """Generate keyword frequency horizontal bar chart from news titles."""
    import re
    from collections import Counter
    if not news:
        return None
    plt = _setup_korean_font()

    all_titles = " ".join(n.get("title", "") for n in news)
    words = re.findall(r"[가-힣]{2,}", all_titles)
    stopwords = {"것으로", "에서", "관련", "대한", "위한", "으로", "이번", "오늘", "내일",
                 "지난", "올해", "이후", "까지", "부터", "하는", "있는", "없는", "되는",
                 "한다", "했다", "라며", "이다", "따라", "통해", "대해"}
    words = [w for w in words if w not in stopwords]
    freq = Counter(words).most_common(12)
    if not freq:
        return None

    labels, counts = zip(*reversed(freq))
    fig, ax = plt.subplots(figsize=(7, max(3, len(labels) * 0.35)))
    colors = ["#1976D2" if i < len(labels) - 3 else "#E65100" for i in range(len(labels))]
    bars = ax.barh(range(len(labels)), counts, color=colors, height=0.6, alpha=0.85)
    for i, (bar, val) in enumerate(zip(bars, counts)):
        ax.text(val + 0.3, i, str(val), va="center", fontsize=8, fontweight="bold")
    ax.set_yticks(range(len(labels)))
    ax.set_yticklabels(labels, fontsize=9)
    ax.set_xlabel("빈도", fontsize=10)
    ax.set_title(f"'{query}' 뉴스 핵심 키워드 빈도 분석", fontsize=11, fontweight="bold")
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    ax.grid(axis="x", alpha=0.2)
    fig.tight_layout()
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=150, bbox_inches="tight")
    plt.close(fig)
    buf.seek(0)
    return buf


def _make_sentiment_pie_chart(news: list) -> io.BytesIO | None:
    """Generate sentiment distribution pie chart."""
    if not news:
        return None
    plt = _setup_korean_font()

    pos_words = {"상승", "급등", "호재", "성장", "개선", "회복", "활황", "강세", "최고",
                 "돌파", "증가", "확대", "호조", "긍정", "수혜", "기대", "추천", "인기",
                 "혁신", "반등", "호황", "흑자"}
    neg_words = {"하락", "급락", "악재", "위축", "감소", "둔화", "약세", "최저",
                 "폭락", "축소", "부진", "우려", "경고", "위기", "리스크", "적자",
                 "침체", "하방", "손실", "붕괴", "악화", "불안", "충격", "규제"}

    pos_count = neg_count = neutral_count = 0
    for n in news:
        title = n.get("title", "")
        has_pos = any(w in title for w in pos_words)
        has_neg = any(w in title for w in neg_words)
        if has_pos and not has_neg:
            pos_count += 1
        elif has_neg and not has_pos:
            neg_count += 1
        elif has_pos and has_neg:
            pos_count += 0.5
            neg_count += 0.5
        else:
            neutral_count += 1

    values = [pos_count, neg_count, neutral_count]
    labels = [f"긍정 ({pos_count:.0f}건)", f"부정 ({neg_count:.0f}건)", f"중립 ({neutral_count:.0f}건)"]
    colors = ["#4CAF50", "#F44336", "#9E9E9E"]
    explode = (0.05, 0.05, 0)

    fig, ax = plt.subplots(figsize=(5, 4))
    wedges, texts, autotexts = ax.pie(values, labels=labels, colors=colors, explode=explode,
                                       autopct="%1.0f%%", startangle=90, pctdistance=0.75)
    for t in autotexts:
        t.set_fontsize(10)
        t.set_fontweight("bold")
    ax.set_title("뉴스 감성 분포 분석", fontsize=11, fontweight="bold", pad=15)
    fig.tight_layout()
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=150, bbox_inches="tight")
    plt.close(fig)
    buf.seek(0)
    return buf


def _make_source_dist_chart(news: list) -> io.BytesIO | None:
    """Generate news source distribution chart."""
    from collections import Counter
    if not news:
        return None
    plt = _setup_korean_font()

    sources = [n.get("source", "기타") or "기타" for n in news]
    freq = Counter(sources).most_common(8)
    if not freq:
        return None

    labels, counts = zip(*freq)
    colors = ["#1976D2", "#388E3C", "#F57C00", "#7B1FA2", "#C62828",
              "#00838F", "#4E342E", "#546E7A"][:len(labels)]

    fig, ax = plt.subplots(figsize=(5, 4))
    wedges, texts, autotexts = ax.pie(counts, labels=labels, colors=colors,
                                       autopct="%1.0f%%", startangle=90, pctdistance=0.8)
    for t in texts:
        t.set_fontsize(8)
    for t in autotexts:
        t.set_fontsize(8)
        t.set_fontweight("bold")
    ax.set_title("뉴스 출처 분포", fontsize=11, fontweight="bold", pad=15)
    fig.tight_layout()
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=150, bbox_inches="tight")
    plt.close(fig)
    buf.seek(0)
    return buf


def _make_daily_change_chart(trend: list, query: str) -> io.BytesIO | None:
    """Generate daily change rate bar chart with outlier detection."""
    if not trend or len(trend) < 3:
        return None
    plt = _setup_korean_font()

    values = [r.get("Trend", 0) for r in trend]
    dates = [str(r.get("Date", "")) for r in trend]
    changes = []
    change_dates = []
    for i in range(1, len(values)):
        if values[i - 1] != 0:
            pct = ((values[i] - values[i - 1]) / values[i - 1]) * 100
        else:
            pct = 0
        changes.append(pct)
        change_dates.append(dates[i])

    if not changes:
        return None

    # Outlier detection (beyond 1.5 * IQR)
    mean_c = sum(changes) / len(changes)
    std_c = (sum((c - mean_c) ** 2 for c in changes) / max(len(changes) - 1, 1)) ** 0.5
    threshold = 1.5 * std_c if std_c > 0 else float("inf")

    colors = []
    for c in changes:
        if abs(c - mean_c) > threshold:
            colors.append("#FF6F00")  # outlier
        elif c >= 0:
            colors.append("#4CAF50")
        else:
            colors.append("#F44336")

    fig, ax = plt.subplots(figsize=(7, 3.5))
    bars = ax.bar(range(len(changes)), changes, color=colors, alpha=0.85, width=0.6)
    for i, (bar, val) in enumerate(zip(bars, changes)):
        y_offset = 0.3 if val >= 0 else -0.3
        ax.text(i, val + y_offset, f"{val:+.1f}%", ha="center", fontsize=8, fontweight="bold",
                color=colors[i])

    ax.axhline(y=0, color="black", linewidth=0.5)
    ax.axhline(y=mean_c, color="#1976D2", linewidth=1, linestyle="--", alpha=0.6,
               label=f"평균: {mean_c:+.1f}%")
    if std_c > 0:
        ax.axhspan(mean_c - threshold, mean_c + threshold, alpha=0.05, color="#1976D2")

    ax.set_xticks(range(len(change_dates)))
    ax.set_xticklabels(change_dates, rotation=30, ha="right", fontsize=9)
    ax.set_ylabel("변화율 (%)", fontsize=10)
    ax.set_title(f"'{query}' 일별 변화율 분석 (주황=이상치)", fontsize=11, fontweight="bold")
    ax.legend(fontsize=8)
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    ax.grid(axis="y", alpha=0.2)
    fig.tight_layout()
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=150, bbox_inches="tight")
    plt.close(fig)
    buf.seek(0)
    return buf


def _make_comparison_bar_chart(context_list) -> io.BytesIO:
    """Generate a cross-domain comparison bar chart for master report."""
    plt = _setup_korean_font()

    names = []
    changes = []
    colors = []

    for item in context_list:
        if not isinstance(item, dict):
            continue
        trend = item.get("df", [])
        if not trend or len(trend) < 2:
            continue
        first = trend[0].get("Trend", 0)
        last = trend[-1].get("Trend", 0)
        if first == 0:
            continue
        pct = ((last - first) / first) * 100
        name = item.get("expert", item.get("query", ""))[:8]
        names.append(name)
        changes.append(pct)
        colors.append("#4CAF50" if pct >= 0 else "#F44336")

    if not names:
        return None

    fig, ax = plt.subplots(figsize=(8, max(3, len(names) * 0.4)))
    bars = ax.barh(range(len(names)), changes, color=colors, height=0.6, alpha=0.85)

    for i, (bar, val) in enumerate(zip(bars, changes)):
        offset = 1 if val >= 0 else -1
        ax.text(val + offset, i, f"{val:+.1f}%", va="center",
                fontsize=8, fontweight="bold",
                color="#4CAF50" if val >= 0 else "#F44336")

    ax.set_yticks(range(len(names)))
    ax.set_yticklabels(names, fontsize=9)
    ax.set_xlabel("변동률 (%)", fontsize=10)
    ax.set_title("분야별 트렌드 변동률 비교", fontsize=12, fontweight="bold")
    ax.axvline(x=0, color="black", linewidth=0.5)
    ax.grid(axis="x", alpha=0.2)
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    fig.tight_layout()

    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=150, bbox_inches="tight")
    plt.close(fig)
    buf.seek(0)
    return buf


# ═══════════════════════════════════════════════════════════════════════════════
# Text cleaning helpers
# ═══════════════════════════════════════════════════════════════════════════════

def _clean_text(text: str) -> str:
    """Strip HTML tags, decode entities, clean whitespace for report output."""
    import re
    from html import unescape
    if not text:
        return ""
    clean = re.sub(r"<[^>]+>", " ", text)
    clean = unescape(clean)
    clean = re.sub(r"\s+", " ", clean).strip()
    # Remove raw URLs masquerading as summaries
    if clean.startswith("http://") or clean.startswith("https://"):
        return ""
    return clean


# ═══════════════════════════════════════════════════════════════════════════════
# Word helper: hyperlinks & bookmarks
# ═══════════════════════════════════════════════════════════════════════════════

def _add_hyperlink(paragraph, text, url):
    """Add a clickable hyperlink to a python-docx paragraph."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    part = paragraph.part
    r_id = part.relate_to(url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    rPr.append(color)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "18")
    rPr.append(sz)
    new_run.append(rPr)

    t_elem = OxmlElement("w:t")
    t_elem.set(qn("xml:space"), "preserve")
    t_elem.text = text
    new_run.append(t_elem)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return paragraph


def _add_bookmark(paragraph, bookmark_name):
    """Add a bookmark anchor to a paragraph for internal linking."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import random

    bm_id = str(random.randint(1000, 99999))
    bm_start = OxmlElement("w:bookmarkStart")
    bm_start.set(qn("w:id"), bm_id)
    bm_start.set(qn("w:name"), bookmark_name)
    paragraph._p.insert(0, bm_start)

    bm_end = OxmlElement("w:bookmarkEnd")
    bm_end.set(qn("w:id"), bm_id)
    paragraph._p.append(bm_end)


def _add_internal_link(paragraph, text, bookmark_name):
    """Add an internal hyperlink (bookmark ref) to a paragraph."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), bookmark_name)

    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "1565C0")
    rPr.append(color)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "20")
    rPr.append(sz)
    new_run.append(rPr)

    t_elem = OxmlElement("w:t")
    t_elem.set(qn("xml:space"), "preserve")
    t_elem.text = text
    new_run.append(t_elem)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def _add_toc_entry(doc, num, title_text, bookmark_name):
    """Add a TOC entry with number, dotted leader, and internal link."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    p = doc.add_paragraph()
    # Set tab stop with dot leader at right margin
    pPr = p._p.get_or_add_pPr()
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "right")
    tab.set(qn("w:leader"), "dot")
    tab.set(qn("w:pos"), "9072")  # ~16cm right margin
    tabs.append(tab)
    pPr.append(tabs)

    # Section number + title as internal link
    _add_internal_link(p, f"{num}. {title_text}", bookmark_name)

    # Tab + page indicator
    tab_run = OxmlElement("w:r")
    tab_char = OxmlElement("w:tab")
    tab_run.append(tab_char)
    p._p.append(tab_run)

    # Page number placeholder
    pg_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "18")
    rPr.append(sz)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "888888")
    rPr.append(color)
    pg_run.append(rPr)
    t = OxmlElement("w:t")
    t.text = f"p.{num}"
    pg_run.append(t)
    p._p.append(pg_run)

    return p


# ═══════════════════════════════════════════════════════════════════════════════
# Individual Expert Word Report
# ═══════════════════════════════════════════════════════════════════════════════

def _gen_word(query, news, web, trend, now_str, table_data=None, youtube=None) -> bytes:
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
    except ImportError:
        return _gen_text(query, news, web, trend, now_str, table_data=table_data, youtube=youtube)

    domain = _match_expert_domain(query)
    section_plan = _plan_individual_sections(query, news, web, trend, table_data, youtube)
    doc = Document()
    _set_update_fields_on_open(doc)

    # ── Title Page ──
    for _ in range(3):
        doc.add_paragraph("")
    title = doc.add_heading(f"{domain['icon']} 생활정보 심층 분석 리포트", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_bookmark(title, "top")

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(f"분석 키워드: {query}")
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x1A, 0x23, 0x7E)
    run.bold = True

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run(f"분석 프레임워크: {domain['framework']}")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    meta2 = doc.add_paragraph()
    meta2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta2.add_run(f"생성일시: {now_str}  |  자동 생성 리포트")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    doc.add_page_break()

    # ── 목차 (표준 Word TOC 필드 — 헤딩 기반 실제 페이지번호 자동 생성) ──
    _toc_label = doc.add_paragraph()
    _tlr = _toc_label.add_run("목차 (Table of Contents)")
    _tlr.bold = True
    _tlr.font.size = Pt(15)
    _tlr.font.color.rgb = RGBColor(0x1A, 0x23, 0x7E)
    _add_standard_toc(doc, [(str(n), t, b) for n, t, b in section_plan], Pt, RGBColor)
    doc.add_paragraph("")

    # ── Disclaimer ──
    if domain.get("disclaimer"):
        p = doc.add_paragraph()
        run = p.add_run(f"⚠ {domain['disclaimer']}")
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0xAA, 0x55, 0x00)
        run.italic = True

    doc.add_page_break()

    # Track section numbers for sequential numbering
    current_sec = 1

    # ══ 1. Executive Summary ══
    h = doc.add_heading(f"{current_sec}. 분석 개요 (Executive Summary)", level=2)
    _add_bookmark(h, "sec_summary")
    current_sec += 1

    total_news = len(news) if news else 0
    total_web = len(web) if web else 0

    summary_parts = [
        f"본 리포트는 '{query}' 키워드를 기반으로 최신 데이터를 수집·분석한 결과입니다.",
        f"총 {total_news}건의 뉴스와 {total_web}건의 웹 검색 결과를 수집하여 다각도로 분석하였습니다.",
    ]

    stats = None
    if trend and len(trend) >= 2:
        values = [r.get("Trend", 0) for r in trend]
        stats = _calc_statistics(values)
        summary_parts.append(
            f"최근 {len(trend)}일간 트렌드 지수는 {stats['first']:,.0f} → {stats['last']:,.0f}으로 "
            f"{abs(stats['pct_change']):.1f}% {stats['trend_dir']}하였으며, "
            f"변동성은 {stats['volatility']} 수준(CV={stats['cv']:.1f}%)입니다."
        )

    # Table data summary for non-trend pages (oil/exchange/stock)
    if table_data:
        summary_parts.append(
            f"총 {len(table_data)}건의 실시간 데이터를 수집하여 현황을 분석하였습니다."
        )
        # Extract numeric values for quick overview
        _numeric_cols = {}
        for row in table_data:
            for k, v in row.items():
                if isinstance(v, (int, float)):
                    _numeric_cols.setdefault(k, []).append(v)
        if _numeric_cols:
            highlights = []
            for col_name, vals in list(_numeric_cols.items())[:3]:
                avg_v = sum(vals) / len(vals)
                highlights.append(f"{col_name} 평균 {avg_v:,.2f}")
            if highlights:
                summary_parts.append(f"주요 수치: {' | '.join(highlights)}")

    # Brief news summary (detailed flow analysis in dedicated section)
    if news:
        themes = _extract_news_themes(news)
        if themes:
            summary_parts.append(
                f"뉴스 핵심 키워드: 【{'、'.join(themes[:5])}】 — 상세 흐름 분석은 별도 섹션을 참조하세요."
            )

    for part in summary_parts:
        p = doc.add_paragraph(part)
        p.paragraph_format.space_after = Pt(4)

    # Key metrics table
    if domain.get("metrics"):
        doc.add_paragraph("")
        p = doc.add_paragraph()
        run = p.add_run("▶ 주요 분석 지표:")
        run.bold = True
        run.font.size = Pt(10)
        for m in domain["metrics"]:
            doc.add_paragraph(f"  • {m}", style="List Bullet")

    # ══ 검색 주제 심층 분석 (질문 맞춤 코너 — 기본 포맷 보존 + 분야 전문 분석은 데이터 유무와 무관하게 항상) ══
    if query:
        doc.add_paragraph("")
        h = doc.add_heading(f"{current_sec}. 검색 주제 심층 분석", level=2)
        _add_bookmark(h, "sec_focus")
        current_sec += 1
        _add_query_focus_section_word(doc, query, news, web, Pt, RGBColor, domain=domain, stats=stats)

    # ══ 2. Trend Analysis ══
    if trend:
        doc.add_paragraph("")
        h = doc.add_heading(f"{current_sec}. 트렌드 분석 및 통계", level=2)
        _add_bookmark(h, "sec_trend")
        current_sec += 1

        doc.add_paragraph(
            f"아래 차트는 '{query}' 관련 최근 추이를 나타냅니다. "
            f"트렌드 지수는 검색 빈도, 뉴스 노출량, 소셜 미디어 언급량 등을 종합 산출한 복합 지수입니다."
        )

        # Chart
        try:
            chart_buf = _make_trend_chart(trend, query, stats)
            doc.add_picture(chart_buf, width=Inches(5.8))
            last_p = doc.paragraphs[-1]
            last_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception as e:
            doc.add_paragraph(f"[차트 생성 실패: {e}]")

        # Statistics table
        if stats:
            doc.add_paragraph("")
            p = doc.add_paragraph()
            run = p.add_run("▶ 통계 분석 요약")
            run.bold = True
            run.font.size = Pt(11)

            table = doc.add_table(rows=8, cols=2, style="Light Shading Accent 1")
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            stat_data = [
                ("평균 (Mean)", f"{stats['mean']:,.1f}"),
                ("표준편차 (Std Dev)", f"{stats['std_dev']:,.2f}"),
                ("최고점 (High)", f"{stats['max']:,.0f}"),
                ("최저점 (Low)", f"{stats['min']:,.0f}"),
                ("변동폭 (Range)", f"{stats['range']:,.0f}"),
                ("변동계수 (CV)", f"{stats['cv']:.1f}%"),
                ("추세 방향", f"{stats['trend_dir']} (기울기: {stats['slope']:+.2f})"),
                ("기간 변동률", f"{stats['pct_change']:+.1f}%"),
            ]
            for row_idx, (label, val) in enumerate(stat_data):
                table.rows[row_idx].cells[0].text = label
                table.rows[row_idx].cells[1].text = val

            # Trend interpretation
            doc.add_paragraph("")
            p = doc.add_paragraph()
            run = p.add_run("▶ 추세 해석")
            run.bold = True

            values = [r.get("Trend", 0) for r in trend]
            max_val = max(values)
            min_val = min(values)
            max_date = trend[values.index(max_val)].get("Date", "")
            min_date = trend[values.index(min_val)].get("Date", "")

            interp = (
                f"분석 기간 중 최고점은 {max_date} ({max_val:,.0f}), "
                f"최저점은 {min_date} ({min_val:,.0f})이며, "
                f"평균 지수는 {stats['mean']:,.1f}입니다. "
            )
            if stats['cv'] > 15:
                interp += (
                    f"변동계수 {stats['cv']:.1f}%로 높은 변동성을 보이고 있어, "
                    f"관련 동향을 면밀히 주시할 필요가 있습니다. "
                    f"급격한 변동은 외부 이벤트(정책 발표, 글로벌 이슈 등)에 기인했을 가능성이 높으며, "
                    f"단기적 과매도/과매수 구간을 식별하는 것이 중요합니다."
                )
            elif stats['cv'] > 5:
                interp += (
                    f"변동계수 {stats['cv']:.1f}%로 보통 수준의 변동성을 나타내고 있습니다. "
                    f"전반적으로 {stats['trend_dir']} 추세이나, 방향성 전환 가능성을 염두에 두어야 합니다."
                )
            else:
                interp += (
                    f"변동계수 {stats['cv']:.1f}%로 안정적인 추세를 유지하고 있습니다. "
                    f"현재의 {stats['trend_dir']} 기조가 단기적으로 지속될 가능성이 높습니다."
                )
            doc.add_paragraph(interp)

    # ══ Daily Change Rate Analysis ══
    if trend and len(trend) >= 3:
        doc.add_paragraph("")
        h = doc.add_heading(f"{current_sec}. 변화율 분석 및 이상치 탐지", level=2)
        _add_bookmark(h, "sec_change_rate")
        current_sec += 1

        values = [r.get("Trend", 0) for r in trend]
        daily_changes = []
        for idx_c in range(1, len(values)):
            if values[idx_c - 1] != 0:
                pct = ((values[idx_c] - values[idx_c - 1]) / values[idx_c - 1]) * 100
            else:
                pct = 0
            daily_changes.append(pct)

        doc.add_paragraph(
            "일별 변화율을 분석하여 급격한 변동(이상치)을 탐지합니다. "
            "이상치는 1.5×표준편차를 초과하는 변화율로 정의됩니다."
        )

        # Chart
        try:
            change_buf = _make_daily_change_chart(trend, query)
            if change_buf:
                doc.add_picture(change_buf, width=Inches(5.8))
                last_p = doc.paragraphs[-1]
                last_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception:
            pass

        # Change rate stats
        if daily_changes:
            mean_dc = sum(daily_changes) / len(daily_changes)
            std_dc = (sum((c - mean_dc) ** 2 for c in daily_changes) / max(len(daily_changes) - 1, 1)) ** 0.5
            max_dc = max(daily_changes)
            min_dc = min(daily_changes)

            doc.add_paragraph("")
            cr_table = doc.add_table(rows=5, cols=2, style="Light Shading Accent 1")
            cr_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            cr_data = [
                ("평균 일별 변화율", f"{mean_dc:+.2f}%"),
                ("변화율 표준편차", f"{std_dc:.2f}%"),
                ("최대 상승 변화", f"{max_dc:+.2f}%"),
                ("최대 하락 변화", f"{min_dc:+.2f}%"),
                ("이상치 탐지 기준", f"±{1.5 * std_dc:.2f}%"),
            ]
            for row_idx, (label, val) in enumerate(cr_data):
                cr_table.rows[row_idx].cells[0].text = label
                cr_table.rows[row_idx].cells[1].text = val

            # Outlier detection
            threshold = 1.5 * std_dc if std_dc > 0 else float("inf")
            outliers = [(i, c) for i, c in enumerate(daily_changes) if abs(c - mean_dc) > threshold]
            if outliers:
                doc.add_paragraph("")
                p = doc.add_paragraph()
                run = p.add_run(f"▶ 이상치 {len(outliers)}건 탐지:")
                run.bold = True
                for oi, oc in outliers:
                    date_label = trend[oi + 1].get("Date", "") if oi + 1 < len(trend) else ""
                    doc.add_paragraph(f"  ⚠ {date_label}: {oc:+.2f}% (기준 초과)")
            else:
                doc.add_paragraph("분석 기간 내 이상치는 탐지되지 않았습니다. 안정적 변동 패턴입니다.")

    # ══ Table Data Section (for non-trend pages: oil/exchange/stock) ══
    if table_data:
        doc.add_paragraph("")
        h = doc.add_heading(f"{current_sec}. 데이터 현황 및 통계 분석", level=2)
        _add_bookmark(h, "sec_table_data")
        current_sec += 1

        doc.add_paragraph(
            f"'{query}' 관련 실시간 수집 데이터 {len(table_data)}건의 현황입니다."
        )

        # Render data as Word table
        if table_data:
            all_keys = []
            for row in table_data:
                for k in row.keys():
                    if k not in all_keys:
                        all_keys.append(k)

            tbl = doc.add_table(rows=1 + len(table_data), cols=len(all_keys),
                                style="Light Shading Accent 1")
            tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
            # Header
            for ci, key in enumerate(all_keys):
                cell = tbl.rows[0].cells[ci]
                cell.text = str(key)
                for run in cell.paragraphs[0].runs:
                    run.bold = True
            # Data rows
            for ri, row in enumerate(table_data):
                for ci, key in enumerate(all_keys):
                    val = row.get(key, "")
                    tbl.rows[ri + 1].cells[ci].text = str(val)

        # Numeric column statistics
        numeric_cols = {}
        for row in table_data:
            for k, v in row.items():
                if isinstance(v, (int, float)):
                    numeric_cols.setdefault(k, []).append(v)

        if numeric_cols:
            doc.add_paragraph("")
            p = doc.add_paragraph()
            run = p.add_run("▶ 수치 데이터 통계 분석")
            run.bold = True
            run.font.size = Pt(11)

            for col_name, vals in numeric_cols.items():
                if len(vals) < 2:
                    continue
                col_stats = _calc_statistics(vals)
                doc.add_paragraph("")
                p = doc.add_paragraph()
                run = p.add_run(f"📊 {col_name}")
                run.bold = True

                stat_tbl = doc.add_table(rows=6, cols=2, style="Light Shading Accent 1")
                stat_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
                stat_items = [
                    ("항목 수", f"{len(vals)}"),
                    ("평균", f"{col_stats['mean']:,.2f}"),
                    ("표준편차", f"{col_stats['std_dev']:,.2f}"),
                    ("최댓값", f"{col_stats['max']:,.2f}"),
                    ("최솟값", f"{col_stats['min']:,.2f}"),
                    ("변동계수 (CV)", f"{col_stats['cv']:.1f}%"),
                ]
                for row_idx, (label, val) in enumerate(stat_items):
                    stat_tbl.rows[row_idx].cells[0].text = label
                    stat_tbl.rows[row_idx].cells[1].text = val

            # Generate bar chart for first numeric column with labels
            try:
                _setup_korean_font()
                import matplotlib.pyplot as plt

                # Find a label column (first non-numeric column)
                label_key = None
                for k in all_keys:
                    sample_vals = [row.get(k) for row in table_data if row.get(k)]
                    if sample_vals and not isinstance(sample_vals[0], (int, float)):
                        label_key = k
                        break

                # Chart the first numeric column
                first_num_col = list(numeric_cols.keys())[0]
                labels = [str(row.get(label_key, i)) for i, row in enumerate(table_data)] if label_key else [str(i + 1) for i in range(len(table_data))]
                chart_vals = [row.get(first_num_col, 0) for row in table_data]
                # Filter to only rows that have numeric value for this column
                paired = [(l, v) for l, v in zip(labels, chart_vals) if isinstance(v, (int, float))]
                if paired:
                    labels, chart_vals = zip(*paired)
                    fig, ax = plt.subplots(figsize=(10, max(4, len(paired) * 0.5)))
                    colors = ['#1A237E' if v >= 0 else '#C62828' for v in chart_vals]
                    bars = ax.barh(range(len(labels)), chart_vals, color=colors, edgecolor='white', height=0.6)
                    ax.set_yticks(range(len(labels)))
                    ax.set_yticklabels(labels, fontsize=9)
                    ax.set_xlabel(first_num_col, fontsize=10)
                    ax.set_title(f"'{query}' — {first_num_col} 비교", fontsize=12, fontweight="bold")
                    ax.invert_yaxis()
                    for bar, val in zip(bars, chart_vals):
                        ax.text(bar.get_width(), bar.get_y() + bar.get_height() / 2,
                                f" {val:,.2f}", va='center', fontsize=8)
                    plt.tight_layout()
                    chart_buf = io.BytesIO()
                    fig.savefig(chart_buf, format="png", dpi=150, bbox_inches="tight")
                    plt.close(fig)
                    chart_buf.seek(0)
                    doc.add_paragraph("")
                    doc.add_picture(chart_buf, width=Inches(5.8))
                    last_p = doc.paragraphs[-1]
                    last_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception:
                pass

    # ══ News Deep Analysis (Sentiment + Keywords + Sources) ══
    if news:
        doc.add_paragraph("")
        h = doc.add_heading(f"{current_sec}. 뉴스 심층 분석 (감성·키워드·출처)", level=2)
        _add_bookmark(h, "sec_news_analysis")
        current_sec += 1

        doc.add_paragraph(
            f"수집된 {len(news)}건의 뉴스에 대해 감성 분석, 핵심 키워드 빈도, 출처 분포를 "
            f"통계적으로 분석하였습니다."
        )

        # 1) Sentiment pie chart
        doc.add_paragraph("")
        p = doc.add_paragraph()
        run = p.add_run("▶ 감성 분석 (Sentiment Analysis)")
        run.bold = True
        run.font.size = Pt(11)

        try:
            sent_buf = _make_sentiment_pie_chart(news)
            if sent_buf:
                doc.add_picture(sent_buf, width=Inches(4.0))
                last_p = doc.paragraphs[-1]
                last_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception:
            pass

        flow_text = _news_flow_summary(news, query)
        if flow_text:
            doc.add_paragraph(flow_text)

        # 2) Keyword frequency chart
        doc.add_paragraph("")
        p = doc.add_paragraph()
        run = p.add_run("▶ 핵심 키워드 빈도 분석")
        run.bold = True
        run.font.size = Pt(11)

        try:
            kw_buf = _make_keyword_freq_chart(news, query)
            if kw_buf:
                doc.add_picture(kw_buf, width=Inches(5.5))
                last_p = doc.paragraphs[-1]
                last_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception:
            pass

        themes = _extract_news_themes(news)
        if themes:
            doc.add_paragraph(
                f"상위 키워드: {'、'.join(themes[:8])} — "
                f"이들 키워드의 빈도 변화를 추적하면 향후 트렌드 선행 지표로 활용 가능합니다."
            )

        # 3) Source distribution chart
        doc.add_paragraph("")
        p = doc.add_paragraph()
        run = p.add_run("▶ 뉴스 출처 분포 분석")
        run.bold = True
        run.font.size = Pt(11)

        try:
            src_buf = _make_source_dist_chart(news)
            if src_buf:
                doc.add_picture(src_buf, width=Inches(4.0))
                last_p = doc.paragraphs[-1]
                last_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception:
            pass

        from collections import Counter as _Counter
        sources = [n.get("source", "기타") or "기타" for n in news]
        src_freq = _Counter(sources).most_common(5)
        if src_freq:
            src_text = ", ".join(f"{s} ({c}건)" for s, c in src_freq)
            doc.add_paragraph(f"주요 보도 매체: {src_text}")

        # ══ News List (compact, no duplication) ══
        doc.add_paragraph("")
        h = doc.add_heading(f"{current_sec}. 주요 뉴스 목록", level=2)
        _add_bookmark(h, "sec_news_list")
        current_sec += 1

        n_display = min(len(news), 10)
        table = doc.add_table(rows=n_display + 1, cols=4, style="Light Shading Accent 1")
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        headers = ["No.", "제목", "출처", "요약"]
        for j, header in enumerate(headers):
            cell = table.rows[0].cells[j]
            cell.text = header
            for p in cell.paragraphs:
                for run in p.runs:
                    run.bold = True

        for idx, n in enumerate(news[:n_display]):
            row = table.rows[idx + 1]
            row.cells[0].text = str(idx + 1)
            title_text = _clean_text(n.get("title", ""))
            row.cells[1].text = title_text[:60]
            row.cells[2].text = _clean_text(n.get("source", ""))
            snippet = _clean_text(n.get("snippet", ""))
            # Prevent snippet ≈ title duplication (word-overlap check)
            if snippet and _is_text_similar(snippet, title_text):
                snippet = ""
            row.cells[3].text = snippet[:100] + ("..." if len(snippet) > 100 else "") if snippet else "원문 참조"

        # Hyperlinks below table (compact)
        doc.add_paragraph("")
        p = doc.add_paragraph()
        run = p.add_run("▶ 원문 링크:")
        run.bold = True
        run.font.size = Pt(9)
        for i, n in enumerate(news[:n_display]):
            link = n.get("link", "")
            title_text = _clean_text(n.get("title", ""))
            if link:
                link_p = doc.add_paragraph()
                run = link_p.add_run(f"  [{i + 1}] ")
                run.font.size = Pt(8)
                _add_hyperlink(link_p, title_text[:70] + ("..." if len(title_text) > 70 else ""), link)

    # ══ Web Results ══
    if web:
        doc.add_paragraph("")
        h = doc.add_heading(f"{current_sec}. 웹 검색 결과 분석", level=2)
        _add_bookmark(h, "sec_web")
        current_sec += 1

        doc.add_paragraph(
            f"'{query}' 관련 웹 검색 결과 {len(web)}건을 수집하였습니다. "
            f"블로그, 포럼, 전문 사이트 등 다양한 소스에서 수집된 정보입니다."
        )

        for i, w in enumerate(web[:10]):
            title_text = _clean_text(w.get("title", "제목 없음"))
            link = w.get("link", "")
            snippet = _clean_text(w.get("snippet", ""))

            p = doc.add_paragraph()
            run = p.add_run(f"[{i + 1}] {title_text}")
            run.bold = True
            run.font.size = Pt(10)

            if snippet:
                p = doc.add_paragraph()
                run = p.add_run(snippet)
                run.font.size = Pt(9)

            if link:
                link_p = doc.add_paragraph()
                run = link_p.add_run("원문: ")
                run.font.size = Pt(9)
                _add_hyperlink(link_p, title_text[:60] + ("..." if len(title_text) > 60 else ""), link)
            doc.add_paragraph("")

    # ══ YouTube Videos ══
    if youtube:
        doc.add_paragraph("")
        h = doc.add_heading(f"{current_sec}. 관련 YouTube 영상", level=2)
        _add_bookmark(h, "sec_youtube")
        current_sec += 1

        doc.add_paragraph(
            f"'{query}' 관련 YouTube 영상 {len(youtube)}건을 수집하였습니다. "
            f"각 영상의 하이퍼링크를 클릭하면 YouTube에서 시청할 수 있습니다."
        )

        yt_table = doc.add_table(rows=1 + len(youtube), cols=4, style="Light Shading Accent 1")
        yt_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for ci, hdr in enumerate(["제목", "채널", "재생시간", "조회수"]):
            cell = yt_table.rows[0].cells[ci]
            cell.text = hdr
            for run in cell.paragraphs[0].runs:
                run.bold = True
        for yi, yt in enumerate(youtube):
            row = yt_table.rows[yi + 1]
            row.cells[0].text = _clean_text(yt.get("title", ""))[:60]
            row.cells[1].text = _clean_text(yt.get("uploader", ""))
            row.cells[2].text = str(yt.get("duration", ""))
            row.cells[3].text = str(yt.get("view_count", ""))

        # YouTube links
        doc.add_paragraph("")
        p = doc.add_paragraph()
        run = p.add_run("▶ YouTube 링크:")
        run.bold = True
        run.font.size = Pt(9)
        for yi, yt in enumerate(youtube):
            url = yt.get("url", "")
            title_text = _clean_text(yt.get("title", ""))
            if url:
                lp = doc.add_paragraph()
                run = lp.add_run(f"  [{yi + 1}] ")
                run.font.size = Pt(8)
                _add_hyperlink(lp, title_text[:70], url)

    # ══ Expert Comprehensive Insight ══
    doc.add_paragraph("")
    h = doc.add_heading(f"{current_sec}. 전문가 종합 인사이트", level=2)
    _add_bookmark(h, "sec_expert")
    current_sec += 1

    p = doc.add_paragraph()
    run = p.add_run(f"▶ 분석 프레임워크: {domain['framework']}")
    run.bold = True
    run.font.size = Pt(11)

    # Unique forward-looking analysis (not repeating trend/news sections)
    doc.add_paragraph(
        f"이상의 트렌드·뉴스·웹 분석 결과를 {domain['framework']} 관점에서 종합하면 다음과 같습니다."
    )

    insight_parts = []
    # Risk-opportunity matrix based on combined signals
    if stats:
        vol_label = stats["volatility"]
        trend_label = stats["trend_dir"]
        if trend_label == "상승" and vol_label == "낮음":
            insight_parts.append(
                "【기회 우위】 안정적 상승세로 진입 기회가 유효합니다. "
                "다만 외부 충격(정책 전환, 글로벌 이벤트)에 따른 급반전 가능성을 "
                "헤지 전략으로 대비하시기 바랍니다."
            )
        elif trend_label == "상승" and vol_label != "낮음":
            insight_parts.append(
                "【주의 필요】 상승세이나 변동성이 높아 과열 신호를 동반합니다. "
                f"3일 이동평균과의 괴리(현재 지수 대비)를 모니터링하고, "
                f"변동계수 {stats['cv']:.1f}%가 정상 범위로 회귀하는지 점검하세요."
            )
        elif trend_label == "하락":
            insight_parts.append(
                "【리스크 경계】 하락 추세에 있으므로 방어적 포지션을 권장합니다. "
                "반등 시그널(일별 변화율 양전환 2일 연속, 긍정 뉴스 비중 확대)을 "
                "기다린 후 전략 수정을 검토하세요."
            )
        else:
            insight_parts.append(
                "【관망 구간】 방향성이 불분명한 보합 상태입니다. "
                "정책 발표, 실적 시즌, 계절 요인 등 촉매(catalyst)에 따라 "
                "방향이 결정될 가능성이 높으므로 이벤트 일정을 주시하세요."
            )
    elif table_data:
        # For non-trend pages: analyze table data numeric patterns
        numeric_cols = {}
        for row in table_data:
            for k, v in row.items():
                if isinstance(v, (int, float)):
                    numeric_cols.setdefault(k, []).append(v)
        if numeric_cols:
            # Find columns with notable variation
            for col_name, vals in numeric_cols.items():
                if len(vals) >= 2:
                    col_stats = _calc_statistics(vals)
                    if col_stats['cv'] > 15:
                        insight_parts.append(
                            f"【{col_name} 변동 주의】 {col_name}의 변동계수가 {col_stats['cv']:.1f}%로 "
                            f"높은 편차를 보이고 있습니다. 평균 {col_stats['mean']:,.2f}, "
                            f"범위 {col_stats['min']:,.2f}~{col_stats['max']:,.2f}으로 "
                            f"항목 간 편차가 큰 만큼 개별 항목의 추이를 면밀히 모니터링해야 합니다."
                        )
                    elif col_stats['cv'] > 5:
                        insight_parts.append(
                            f"【{col_name} 현황】 {col_name}은 평균 {col_stats['mean']:,.2f}으로, "
                            f"보통 수준의 변동성(CV {col_stats['cv']:.1f}%)을 보이고 있습니다. "
                            f"최대 {col_stats['max']:,.2f}, 최소 {col_stats['min']:,.2f}입니다."
                        )
                    else:
                        insight_parts.append(
                            f"【{col_name} 안정】 {col_name}은 평균 {col_stats['mean']:,.2f}으로 안정적입니다. "
                            f"변동계수 {col_stats['cv']:.1f}%로 항목 간 편차가 적습니다."
                        )
        if not insight_parts:
            insight_parts.append(
                f"수집된 {len(table_data)}건의 데이터를 기반으로 현재 시장 상황을 점검하세요. "
                f"주요 지표의 변화를 정기적으로 모니터링하여 조기 대응 체계를 갖추는 것이 중요합니다."
            )

    # Cross-signal consistency check (news sentiment analysis)
    if news:
        pos_kw = {"상승", "성장", "개선", "회복", "강세", "호조", "확대", "증가"}
        neg_kw = {"하락", "위축", "감소", "약세", "부진", "우려", "침체", "위기"}
        all_titles = " ".join(n.get("title", "") for n in news)
        p_cnt = sum(1 for w in pos_kw if w in all_titles)
        n_cnt = sum(1 for w in neg_kw if w in all_titles)
        news_dir = "긍정" if p_cnt > n_cnt else ("부정" if n_cnt > p_cnt else "중립")

        if stats:
            trend_dir = stats["trend_dir"]
            if (trend_dir == "상승" and news_dir == "긍정") or (trend_dir == "하락" and news_dir == "부정"):
                insight_parts.append(
                    f"트렌드({trend_dir})와 뉴스 감성({news_dir})이 일치하여 "
                    f"현재 방향성의 신뢰도가 높습니다."
                )
            elif (trend_dir == "상승" and news_dir == "부정") or (trend_dir == "하락" and news_dir == "긍정"):
                insight_parts.append(
                    f"⚠ 트렌드({trend_dir})와 뉴스 감성({news_dir})이 상충합니다. "
                    f"이는 조만간 추세 전환이 일어날 수 있는 선행 신호일 수 있으므로 "
                    f"향후 1~2일간의 변화를 면밀히 관찰하세요."
                )
        else:
            # No trend data — provide standalone news sentiment insight
            sentiment_desc = {
                "긍정": "긍정적 뉴스가 우세하여 시장/환경이 호의적인 분위기입니다. 다만 과잉 낙관에 유의하세요.",
                "부정": "부정적 뉴스가 우세하여 경계가 필요한 시점입니다. 리스크 관리에 집중하세요.",
                "중립": "뉴스 감성이 중립적으로 방향성 판단이 어려운 상황입니다. 추가 데이터를 확인하세요.",
            }
            insight_parts.append(
                f"뉴스 감성 분석 결과: 【{news_dir}】(긍정 {p_cnt}건, 부정 {n_cnt}건) — "
                f"{sentiment_desc.get(news_dir, '')}"
            )

    # Action items
    doc.add_paragraph("")
    p = doc.add_paragraph()
    run = p.add_run("▶ 실행 권고사항 (Action Items)")
    run.bold = True
    run.font.size = Pt(11)

    actions = [
        f"1. 모니터링 지표 — {', '.join(domain.get('metrics', [])[:3])} 등을 주기적으로 점검",
        "2. 뉴스 감성 추이 — 긍정/부정 비율 변화를 주 1회 이상 체크",
    ]
    if stats and stats.get("cv", 0) > 10:
        actions.append("3. 변동성 관리 — 변동계수가 높으므로 분산 투자/접근 전략 검토")
    else:
        actions.append("3. 기회 포착 — 안정적 환경에서의 신규 진입/확대 기회 검토")
    actions.append("4. 외부 변수 — 정책, 글로벌 이슈, 계절 요인에 따른 시나리오별 대응 준비")

    for a in actions:
        doc.add_paragraph(a)

    for part in insight_parts:
        doc.add_paragraph(part)

    # ── 분야별 전략 관점 (도메인 특화) ──
    _render_domain_strategy_word(doc, domain, Pt, RGBColor)

    # ══ References ══
    doc.add_page_break()
    h = doc.add_heading(f"{current_sec}. 참고문헌 (References)", level=2)
    _add_bookmark(h, "sec_refs")

    doc.add_paragraph(
        "본 리포트에 인용된 모든 뉴스 및 웹 검색 결과의 원문 링크입니다. "
        "각 항목의 하이퍼링크를 클릭하면 원문 페이지로 이동합니다."
    )

    ref_idx = 1
    if news:
        p = doc.add_paragraph()
        run = p.add_run("■ 뉴스 출처")
        run.bold = True
        for n in news[:10]:
            link = n.get("link", "")
            title_text = n.get("title", "")
            source = n.get("source", "")
            if link:
                p = doc.add_paragraph()
                run = p.add_run(f"  [{ref_idx}] ")
                run.font.size = Pt(9)
                label = f"{title_text}"
                if source:
                    label += f" — {source}"
                _add_hyperlink(p, label[:120], link)
                ref_idx += 1

    if web:
        doc.add_paragraph("")
        p = doc.add_paragraph()
        run = p.add_run("■ 웹 검색 출처")
        run.bold = True
        for w in web[:10]:
            link = w.get("link", "")
            title_text = w.get("title", "")
            if link:
                p = doc.add_paragraph()
                run = p.add_run(f"  [{ref_idx}] ")
                run.font.size = Pt(9)
                _add_hyperlink(p, title_text[:120], link)
                ref_idx += 1

    # ── Footer ──
    doc.add_paragraph("")
    doc.add_paragraph("")
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("━" * 40)
    run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
    footer2 = doc.add_paragraph()
    footer2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer2.add_run(f"생활정보 분석 플랫폼 자동 생성 리포트  |  {now_str}")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
    footer3 = doc.add_paragraph()
    footer3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer3.add_run("Powered by Life-Info Expert Analysis System")
    run.font.size = Pt(7)
    run.font.color.rgb = RGBColor(0xBB, 0xBB, 0xBB)
    run.italic = True

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ═══════════════════════════════════════════════════════════════════════════════
# Master Report (multi-expert)
# ═══════════════════════════════════════════════════════════════════════════════

def _gen_word_master(context_list, now_str) -> bytes:
    """Generate a comprehensive master Word report with per-expert sections."""
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
    except ImportError:
        return _gen_text_master(context_list, now_str)

    valid_items = [item for item in context_list if isinstance(item, dict)]
    doc = Document()
    _set_update_fields_on_open(doc)

    # ══════════ Title Page ══════════
    for _ in range(4):
        doc.add_paragraph("")
    title = doc.add_heading("전 분야 마스터 리포트", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_bookmark(title, "top")

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("생활정보 전문가 시스템 종합 분석 보고서")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x1A, 0x23, 0x7E)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run(f"생성일시: {now_str}  |  {len(valid_items)}개 분야 종합 분석")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    meta2 = doc.add_paragraph()
    meta2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    total_news = sum(len(item.get("news", [])) for item in valid_items)
    total_web = sum(len(item.get("web", [])) for item in valid_items)
    run = meta2.add_run(f"수집 데이터: 뉴스 {total_news}건  |  웹 검색 {total_web}건")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    doc.add_page_break()

    # ══════════ Table of Contents (표준 TOC 필드 + 캐시 항목) ══════════
    # 목차 라벨은 헤딩이 아닌 문단으로 — Word TOC 자동 갱신 시 목차 자신이 항목에 포함되지 않도록.
    toc_label = doc.add_paragraph()
    _add_bookmark(toc_label, "toc")
    _tlr = toc_label.add_run("목차 (Table of Contents)")
    _tlr.bold = True
    _tlr.font.size = Pt(15)
    _tlr.font.color.rgb = RGBColor(0x1A, 0x23, 0x7E)

    master_entries = [
        ("I", "총괄 분석 개요 (Executive Overview)", "sec_overview"),
        ("II", "분야별 트렌드 비교 차트", "sec_comparison"),
    ]
    for i, item in enumerate(valid_items):
        expert_name = item.get("expert", item.get("query", f"분야 {i + 1}"))
        domain = _match_expert_domain(item.get("query", ""), expert_name)
        master_entries.append((str(i + 1), f"{domain.get('icon', '📊')} {expert_name}", f"expert_{i}"))
    master_entries.append(("★", "참고문헌 (References)", "sec_master_refs"))

    _add_standard_toc(doc, master_entries, Pt, RGBColor)

    doc.add_page_break()

    # ══════════ I. Executive Overview ══════════
    h = doc.add_heading("I. 총괄 분석 개요 (Executive Overview)", level=2)
    _add_bookmark(h, "sec_overview")

    doc.add_paragraph(
        f"본 마스터 리포트는 {len(valid_items)}개 생활정보 분야의 최신 동향을 종합 분석한 결과입니다. "
        f"총 {total_news}건의 뉴스와 {total_web}건의 웹 검색 결과를 기반으로, "
        f"각 분야별 전문가 프레임워크를 적용하여 심층 분석하였습니다."
    )

    # Overview table
    doc.add_paragraph("")
    p = doc.add_paragraph()
    run = p.add_run("▶ 분야별 핵심 지표 요약")
    run.bold = True
    run.font.size = Pt(11)

    table = doc.add_table(rows=len(valid_items) + 1, cols=5, style="Light Shading Accent 1")
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["분야", "뉴스", "웹", "추세", "변동률"]
    for j, header in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = header
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True

    for idx, item in enumerate(valid_items):
        expert_name = item.get("expert", "")[:12]
        news_cnt = len(item.get("news", []))
        web_cnt = len(item.get("web", []))
        trend_data = item.get("df", [])
        trend_dir = "-"
        change_str = "-"
        if trend_data and len(trend_data) >= 2:
            first_v = trend_data[0].get("Trend", 0)
            last_v = trend_data[-1].get("Trend", 0)
            if first_v > 0:
                pct = ((last_v - first_v) / first_v) * 100
                trend_dir = "▲" if pct > 0 else ("▼" if pct < 0 else "━")
                change_str = f"{pct:+.1f}%"

        row = table.rows[idx + 1]
        row.cells[0].text = expert_name
        row.cells[1].text = str(news_cnt)
        row.cells[2].text = str(web_cnt)
        row.cells[3].text = trend_dir
        row.cells[4].text = change_str

    # Top risers / fallers
    doc.add_paragraph("")
    ranked = []
    for item in valid_items:
        trend_data = item.get("df", [])
        if trend_data and len(trend_data) >= 2:
            first_v = trend_data[0].get("Trend", 0)
            last_v = trend_data[-1].get("Trend", 0)
            if first_v > 0:
                pct = ((last_v - first_v) / first_v) * 100
                ranked.append((item.get("expert", ""), pct))

    if ranked:
        ranked.sort(key=lambda x: x[1], reverse=True)
        p = doc.add_paragraph()
        run = p.add_run("▶ 상승 Top 3:")
        run.bold = True
        for name, pct in ranked[:3]:
            doc.add_paragraph(f"  📈 {name}: {pct:+.1f}%")

        p = doc.add_paragraph()
        run = p.add_run("▶ 하락 Top 3:")
        run.bold = True
        for name, pct in ranked[-3:]:
            doc.add_paragraph(f"  📉 {name}: {pct:+.1f}%")

    # ══════════ II. Comparison Chart ══════════
    doc.add_page_break()
    h = doc.add_heading("II. 분야별 트렌드 비교 차트", level=2)
    _add_bookmark(h, "sec_comparison")

    doc.add_paragraph(
        "아래 차트는 전 분야의 최근 트렌드 변동률을 비교한 결과입니다. "
        "양(+)의 값은 상승, 음(-)의 값은 하락을 나타내며, "
        "절대값이 클수록 변동 폭이 큰 것을 의미합니다."
    )

    try:
        comp_buf = _make_comparison_bar_chart(valid_items)
        if comp_buf:
            doc.add_picture(comp_buf, width=Inches(5.8))
            last_p = doc.paragraphs[-1]
            last_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    except Exception as e:
        doc.add_paragraph(f"[비교 차트 생성 실패: {e}]")

    # ══════════ Per-Expert Sections ══════════
    all_refs = []
    for i, item in enumerate(valid_items):
        doc.add_page_break()

        expert_name = item.get("expert", item.get("query", f"분야 {i + 1}"))
        query = item.get("query", expert_name)
        news = item.get("news", [])
        web = item.get("web", [])
        trend = item.get("df", [])
        domain = _match_expert_domain(query, expert_name)

        h = doc.add_heading(f"{i + 1}. {domain.get('icon', '📊')} {expert_name}", level=2)
        _add_bookmark(h, f"expert_{i}")

        # Back to TOC link
        p = doc.add_paragraph()
        _add_internal_link(p, "↑ 목차로 이동", "toc")

        # Framework & metrics
        p = doc.add_paragraph()
        run = p.add_run(f"분석 프레임워크: {domain['framework']}")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        run.italic = True

        # Summary
        total_n = len(news)
        total_w = len(web)
        summary = f"'{query}' 키워드 기반 수집 — 뉴스 {total_n}건, 웹 {total_w}건."
        stats = None
        if trend and len(trend) >= 2:
            values = [r.get("Trend", 0) for r in trend]
            stats = _calc_statistics(values)
            summary += (
                f" 트렌드: {stats['first']:,.0f} → {stats['last']:,.0f} "
                f"({stats['pct_change']:+.1f}% {stats['trend_dir']}, "
                f"변동성: {stats['volatility']})"
            )
        doc.add_paragraph(summary)

        # Chart
        if trend:
            try:
                chart_buf = _make_trend_chart(trend, query, stats)
                doc.add_picture(chart_buf, width=Inches(5.2))
                last_p = doc.paragraphs[-1]
                last_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception as e:
                doc.add_paragraph(f"[차트 생성 실패: {e}]")

            # Brief stat summary
            if stats:
                doc.add_paragraph(
                    f"  평균: {stats['mean']:,.1f}  |  "
                    f"최고: {stats['max']:,.0f}  |  최저: {stats['min']:,.0f}  |  "
                    f"변동폭: {stats['range']:,.0f}  |  CV: {stats['cv']:.1f}%"
                )

        # Expert opinion
        doc.add_paragraph("")
        p = doc.add_paragraph()
        run = p.add_run("▶ 전문가 분석 의견")
        run.bold = True

        expert_text = f"'{expert_name}' 분야는 "
        if stats:
            if stats["trend_dir"] == "상승":
                expert_text += (
                    f"최근 {stats['pct_change']:+.1f}%의 상승세를 기록하고 있습니다. "
                    f"관심도가 증가하고 있어 관련 기회 요인을 주목할 필요가 있습니다."
                )
            elif stats["trend_dir"] == "하락":
                expert_text += (
                    f"최근 {stats['pct_change']:+.1f}%의 하락세를 보이고 있어, "
                    f"하락 원인 분석 및 리스크 관리에 유의해야 합니다."
                )
            else:
                expert_text += "보합 상태로, 방향성 전환 신호를 주시해야 합니다."
        else:
            expert_text += "현재 수집된 데이터 기반으로 동향 모니터링 중입니다."

        doc.add_paragraph(expert_text)

        # ── 분야 전문 심층 분석 (compact) ──
        dd = _domain_dd(domain)
        if dd.get("lens"):
            doc.add_paragraph("")
            p = doc.add_paragraph()
            run = p.add_run("▶ 분야 전문 관점")
            run.bold = True
            for t_, d_ in dd["lens"][:3]:
                bp = doc.add_paragraph(style="List Bullet")
                rt = bp.add_run(f"{t_} — ")
                rt.bold = True
                rt.font.size = Pt(9)
                bp.add_run(d_).font.size = Pt(9)
            _m_terms = _query_key_terms(query)
            _m_figs = _extract_figures_from_items((news or []) + (web or []), limit=5)
            doc.add_paragraph(_domain_narrative(domain, query, _m_terms, _m_figs, stats))
            if dd.get("risks"):
                rp = doc.add_paragraph(f"  리스크: {' · '.join(dd['risks'])}")
                for rn in rp.runs:
                    rn.font.size = Pt(9)
                    rn.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            if dd.get("strategy"):
                sp = doc.add_paragraph(f"  전략: {' / '.join(dd['strategy'])}")
                for rn in sp.runs:
                    rn.font.size = Pt(9)
                    rn.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

        # Disclaimer
        if domain.get("disclaimer"):
            p = doc.add_paragraph()
            run = p.add_run(f"⚠ {domain['disclaimer']}")
            run.font.size = Pt(7)
            run.font.color.rgb = RGBColor(0xAA, 0x55, 0x00)
            run.italic = True

        # News flow
        if news:
            doc.add_paragraph("")
            p = doc.add_paragraph()
            run = p.add_run("▶ 뉴스 흐름 분석")
            run.bold = True

            flow = _news_flow_summary(news, query)
            if flow:
                doc.add_paragraph(flow)

            # News items
            p = doc.add_paragraph()
            run = p.add_run("▶ 주요 뉴스:")
            run.bold = True
            for n in news[:5]:
                title_text = _clean_text(n.get("title", ""))
                link = n.get("link", "")
                source = _clean_text(n.get("source", ""))
                snippet = _clean_text(n.get("snippet", ""))
                if snippet and _is_text_similar(snippet, title_text):
                    snippet = ""

                np_p = doc.add_paragraph()
                run = np_p.add_run(f"• {title_text}")
                run.font.size = Pt(10)
                run.bold = True

                if snippet:
                    sp = doc.add_paragraph()
                    run = sp.add_run(f"  {snippet[:120]}")
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

                if link:
                    lp = doc.add_paragraph()
                    run = lp.add_run("  원문: ")
                    run.font.size = Pt(8)
                    _add_hyperlink(lp, title_text[:60] + ("..." if len(title_text) > 60 else ""), link)
                    all_refs.append({"title": title_text, "source": source, "link": link, "domain": expert_name})

        # Web results
        if web:
            doc.add_paragraph("")
            p = doc.add_paragraph()
            run = p.add_run("▶ 웹 검색 결과:")
            run.bold = True
            for w in web[:3]:
                title_text = _clean_text(w.get("title", ""))
                link = w.get("link", "")
                snippet = _clean_text(w.get("snippet", ""))

                wp = doc.add_paragraph()
                run = wp.add_run(f"• {title_text}")
                run.font.size = Pt(10)

                if snippet:
                    sp = doc.add_paragraph()
                    run = sp.add_run(f"  {snippet[:120]}")
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

                if link:
                    lp = doc.add_paragraph()
                    run = lp.add_run("  원문: ")
                    run.font.size = Pt(8)
                    _add_hyperlink(lp, title_text[:60] + ("..." if len(title_text) > 60 else ""), link)
                    all_refs.append({"title": title_text, "source": "", "link": link, "domain": expert_name})

    # ══════════ References ══════════
    doc.add_page_break()
    h = doc.add_heading("참고문헌 (References)", level=2)
    _add_bookmark(h, "sec_master_refs")

    doc.add_paragraph(
        f"본 마스터 리포트에 인용된 총 {len(all_refs)}건의 원문 링크입니다. "
        f"각 항목의 하이퍼링크를 클릭하면 원문 페이지로 이동합니다."
    )

    # Group by domain
    from collections import defaultdict
    refs_by_domain = defaultdict(list)
    for ref in all_refs:
        refs_by_domain[ref.get("domain", "기타")].append(ref)

    ref_global_idx = 1
    for domain_name, refs in refs_by_domain.items():
        p = doc.add_paragraph()
        run = p.add_run(f"■ {domain_name}")
        run.bold = True
        run.font.size = Pt(10)

        for ref in refs:
            if ref.get("link"):
                p = doc.add_paragraph()
                run = p.add_run(f"  [{ref_global_idx}] ")
                run.font.size = Pt(8)
                label = ref.get("title", "")[:100]
                if ref.get("source"):
                    label += f" — {ref['source']}"
                _add_hyperlink(p, label, ref["link"])
                ref_global_idx += 1

    # ── Footer ──
    doc.add_paragraph("")
    doc.add_paragraph("")
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("━" * 40)
    run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
    footer2 = doc.add_paragraph()
    footer2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer2.add_run(f"전 분야 마스터 리포트 — 자동 생성  |  {now_str}")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
    footer3 = doc.add_paragraph()
    footer3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer3.add_run("Powered by Life-Info Expert Analysis System")
    run.font.size = Pt(7)
    run.font.color.rgb = RGBColor(0xBB, 0xBB, 0xBB)
    run.italic = True

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ═══════════════════════════════════════════════════════════════════════════════
# Plain text generators
# ═══════════════════════════════════════════════════════════════════════════════

def _gen_text(query, news, web, trend, now_str, table_data=None, youtube=None) -> bytes:
    lines = [f"생활정보 분석 리포트 — {query}", f"생성일시: {now_str}", "=" * 60, ""]

    # ── 검색 주제 심층 분석 (질문 맞춤 코너 — 분야 전문 분석은 항상) ──
    if query:
        relevant, terms, n_matched, scope = _rank_relevant_items(query, news, web)
        lines.append(f"[ 검색 주제 심층 분석 — '{query}' ]")
        if terms:
            lines.append("  핵심 검색어: " + ", ".join(terms[:6]))
        if scope:
            lines.append(f"  기간 의도: 최근 {scope} 관점")
        lines.append(f"  직접 매칭 자료: {n_matched}건")
        figs = _extract_figures_from_items(relevant, limit=8)
        _dd_stats = None
        if trend:
            _vals = [r.get("Trend", 0) for r in trend]
            if len(_vals) >= 2:
                _dd_stats = _calc_statistics(_vals)
        _tdomain = _match_expert_domain(query)
        # 주제 핵심 논점 종합 (검색 주제 상세 내용)
        lines.append("  [ 주제 핵심 논점 종합 ]")
        for para in _topic_synthesis_paras(query, news, web, relevant, terms, figs, scope, n_matched, _tdomain, _dd_stats):
            lines.append(f"    {para}")
        if relevant:
            lines.append("  주제 관련 핵심 자료(관련도 순):")
            for it in relevant:
                lines.append(f"    - [{it['_kind']}] {_clean_text(it.get('title', ''))[:120]}")
                snip = _clean_text(it.get("snippet", "") or "")[:160]
                if snip:
                    lines.append(f"      요지: {snip}")
                if it.get("link"):
                    lines.append(f"      링크: {it.get('link')}")
        if figs:
            lines.append("  추출 수치·사실: " + " · ".join(figs))
            lines.append("  ※ 자동 추출 참고값 — 원문 링크로 검증 요망")
        # 분야 전문 심층 분석 (검색 종류 맞춤)
        lines.extend(_domain_deepdive_text_lines(_tdomain, query, _dd_stats, figs, terms))
        lines.append("")

    if trend:
        values = [r.get("Trend", 0) for r in trend]
        stats = _calc_statistics(values) if len(values) >= 2 else {}
        lines.append("[ 트렌드 데이터 ]")
        for row in trend:
            lines.append(f"  {row.get('Date', '')}  →  {row.get('Trend', '')}")
        if stats:
            lines.append(f"\n  평균: {stats['mean']:,.1f}  |  변동률: {stats['pct_change']:+.1f}%  |  추세: {stats['trend_dir']}")
        lines.append("")

    if table_data:
        lines.append("[ 데이터 현황 ]")
        for row in table_data:
            parts = [f"{k}: {v}" for k, v in row.items()]
            lines.append(f"  {' | '.join(parts)}")
        # Numeric stats
        numeric_cols = {}
        for row in table_data:
            for k, v in row.items():
                if isinstance(v, (int, float)):
                    numeric_cols.setdefault(k, []).append(v)
        for col_name, vals in numeric_cols.items():
            if len(vals) >= 2:
                col_stats = _calc_statistics(vals)
                lines.append(f"\n  [{col_name}] 평균: {col_stats['mean']:,.2f}  |  최대: {col_stats['max']:,.2f}  |  최소: {col_stats['min']:,.2f}  |  CV: {col_stats['cv']:.1f}%")
        lines.append("")

    if news:
        lines.append("[ 관련 뉴스 ]")
        for n in news[:10]:
            title_t = n.get('title', '')
            lines.append(f"  - {title_t} ({n.get('source', '')})")
            snippet = n.get('snippet', '')
            if snippet and not _is_text_similar(snippet, title_t):
                lines.append(f"    요약: {snippet[:100]}")
            lines.append(f"    링크: {n.get('link', '')}")
        flow = _news_flow_summary(news, query)
        if flow:
            lines.append(f"\n  [흐름분석] {flow}")
        lines.append("")

    if web:
        lines.append("[ 웹 검색 결과 ]")
        for w in web[:10]:
            lines.append(f"  - {w.get('title', '')}")
            lines.append(f"    {w.get('link', '')}")
        lines.append("")

    if youtube:
        lines.append("[ 관련 YouTube 영상 ]")
        for yt in youtube:
            lines.append(f"  - {yt.get('title', '')} ({yt.get('duration', '')})")
            meta = []
            if yt.get("uploader"):
                meta.append(f"채널: {yt['uploader']}")
            if yt.get("view_count"):
                meta.append(f"조회수: {yt['view_count']}")
            if meta:
                lines.append(f"    {' | '.join(meta)}")
            lines.append(f"    링크: {yt.get('url', '')}")
        lines.append("")

    return "\n".join(lines).encode("utf-8")


def _gen_text_master(context_list, now_str) -> bytes:
    lines = ["전 분야 마스터 리포트", f"생성일시: {now_str}", "=" * 60, ""]
    for i, item in enumerate(context_list):
        if not isinstance(item, dict):
            continue
        expert_name = item.get("expert", item.get("query", f"분야 {i + 1}"))
        query = item.get("query", expert_name)
        lines.append(f"{'─' * 40}")
        lines.append(f"■ [{i + 1}] {expert_name}")
        lines.append(f"  검색어: {query}")
        trend = item.get("df", [])
        if trend:
            values = [r.get("Trend", 0) for r in trend]
            stats = _calc_statistics(values) if len(values) >= 2 else {}
            for row in trend:
                lines.append(f"  {row.get('Date', '')} → {row.get('Trend', '')}")
            if stats:
                lines.append(f"  평균: {stats['mean']:,.1f}  |  변동률: {stats['pct_change']:+.1f}%  |  추세: {stats['trend_dir']}")
        for n in item.get("news", [])[:5]:
            lines.append(f"  뉴스: {n.get('title', '')} ({n.get('source', '')})")
            lines.append(f"        {n.get('link', '')}")
        for w in item.get("web", [])[:3]:
            lines.append(f"  웹: {w.get('title', '')}")
            lines.append(f"      {w.get('link', '')}")
        # 분야 전문 심층 분석
        _m_domain = _match_expert_domain(query, expert_name)
        _m_stats = None
        if trend:
            _m_vals = [r.get("Trend", 0) for r in trend]
            if len(_m_vals) >= 2:
                _m_stats = _calc_statistics(_m_vals)
        _m_rel, _m_terms, _, _ = _rank_relevant_items(query, item.get("news", []), item.get("web", []))
        _m_figs = _extract_figures_from_items(_m_rel, limit=6)
        lines.extend(_domain_deepdive_text_lines(_m_domain, query, _m_stats, _m_figs, _m_terms))
        lines.append("")
    return "\n".join(lines).encode("utf-8")


# ═══════════════════════════════════════════════════════════════════════════════
# Excel generators
# ═══════════════════════════════════════════════════════════════════════════════

def _gen_excel(query, news, web, trend, now_str, table_data=None, youtube=None) -> bytes:
    try:
        import xlsxwriter
        buf = io.BytesIO()
        wb = xlsxwriter.Workbook(buf)
        bold = wb.add_format({"bold": True})
        header_fmt = wb.add_format({"bold": True, "bg_color": "#1A237E", "font_color": "white"})

        if trend:
            # Trend sheet
            ws1 = wb.add_worksheet("트렌드")
            ws1.write(0, 0, f"분석 키워드: {query}", bold)
            ws1.write(1, 0, f"생성일시: {now_str}")
            ws1.write(3, 0, "날짜", header_fmt)
            ws1.write(3, 1, "트렌드", header_fmt)
            for i, row in enumerate(trend or []):
                ws1.write(4 + i, 0, str(row.get("Date", "")))
                ws1.write(4 + i, 1, row.get("Trend", 0))

            # Statistics
            if len(trend) >= 2:
                values = [r.get("Trend", 0) for r in trend]
                stats = _calc_statistics(values)
                r = 4 + len(trend) + 1
                ws1.write(r, 0, "통계", bold)
                stat_items = [("평균", stats['mean']), ("표준편차", stats['std_dev']),
                              ("최고", stats['max']), ("최저", stats['min']),
                              ("변동률(%)", stats['pct_change'])]
                for j, (label, val) in enumerate(stat_items):
                    ws1.write(r + 1 + j, 0, label)
                    ws1.write(r + 1 + j, 1, round(val, 2))

        if table_data:
            # Table data sheet
            ws_data = wb.add_worksheet("데이터 현황")
            ws_data.write(0, 0, f"분석 키워드: {query}", bold)
            ws_data.write(1, 0, f"생성일시: {now_str}")
            all_keys = []
            for row in table_data:
                for k in row.keys():
                    if k not in all_keys:
                        all_keys.append(k)
            for ci, key in enumerate(all_keys):
                ws_data.write(3, ci, key, header_fmt)
                ws_data.set_column(ci, ci, max(15, len(str(key)) * 2))
            for ri, row in enumerate(table_data):
                for ci, key in enumerate(all_keys):
                    val = row.get(key, "")
                    if isinstance(val, (int, float)):
                        ws_data.write(4 + ri, ci, val)
                    else:
                        ws_data.write(4 + ri, ci, str(val))

            # Statistics section below data
            numeric_cols = {}
            for row in table_data:
                for k, v in row.items():
                    if isinstance(v, (int, float)):
                        numeric_cols.setdefault(k, []).append(v)
            if numeric_cols:
                stat_row = 4 + len(table_data) + 2
                ws_data.write(stat_row, 0, "통계 분석", bold)
                stat_row += 1
                ws_data.write(stat_row, 0, "항목", header_fmt)
                ws_data.write(stat_row, 1, "평균", header_fmt)
                ws_data.write(stat_row, 2, "표준편차", header_fmt)
                ws_data.write(stat_row, 3, "최대", header_fmt)
                ws_data.write(stat_row, 4, "최소", header_fmt)
                ws_data.write(stat_row, 5, "CV(%)", header_fmt)
                for ci, (col_name, vals) in enumerate(numeric_cols.items()):
                    if len(vals) >= 2:
                        cs = _calc_statistics(vals)
                        r = stat_row + 1 + ci
                        ws_data.write(r, 0, col_name)
                        ws_data.write(r, 1, round(cs['mean'], 2))
                        ws_data.write(r, 2, round(cs['std_dev'], 2))
                        ws_data.write(r, 3, round(cs['max'], 2))
                        ws_data.write(r, 4, round(cs['min'], 2))
                        ws_data.write(r, 5, round(cs['cv'], 1))

        # News sheet
        ws2 = wb.add_worksheet("뉴스")
        ws2.write(0, 0, "No.", header_fmt)
        ws2.write(0, 1, "제목", header_fmt)
        ws2.write(0, 2, "출처", header_fmt)
        ws2.write(0, 3, "요약", header_fmt)
        ws2.write(0, 4, "링크", header_fmt)
        for i, n in enumerate(news or []):
            _title = n.get("title", "")
            _snip = n.get("snippet", "")
            if _snip and _is_text_similar(_snip, _title):
                _snip = ""
            ws2.write(1 + i, 0, i + 1)
            ws2.write(1 + i, 1, _title)
            ws2.write(1 + i, 2, n.get("source", ""))
            ws2.write(1 + i, 3, _snip or "원문 참조")
            ws2.write(1 + i, 4, n.get("link", ""))
        ws2.set_column(1, 1, 40)
        ws2.set_column(3, 3, 50)
        ws2.set_column(4, 4, 60)

        # YouTube sheet
        if youtube:
            ws3 = wb.add_worksheet("YouTube")
            ws3.write(0, 0, "No.", header_fmt)
            ws3.write(0, 1, "제목", header_fmt)
            ws3.write(0, 2, "채널", header_fmt)
            ws3.write(0, 3, "재생시간", header_fmt)
            ws3.write(0, 4, "조회수", header_fmt)
            ws3.write(0, 5, "링크", header_fmt)
            for i, yt in enumerate(youtube):
                ws3.write(1 + i, 0, i + 1)
                ws3.write(1 + i, 1, yt.get("title", ""))
                ws3.write(1 + i, 2, yt.get("uploader", ""))
                ws3.write(1 + i, 3, yt.get("duration", ""))
                ws3.write(1 + i, 4, yt.get("view_count", ""))
                ws3.write(1 + i, 5, yt.get("url", ""))
            ws3.set_column(1, 1, 40)
            ws3.set_column(5, 5, 60)

        wb.close()
        return buf.getvalue()
    except ImportError:
        from openpyxl import Workbook
        wb = Workbook()
        ws = wb.active
        if trend:
            ws.title = "트렌드"
            ws.append([f"분석 키워드: {query}", f"생성일시: {now_str}"])
            ws.append(["날짜", "트렌드"])
            for row in trend:
                ws.append([str(row.get("Date", "")), row.get("Trend", 0)])
        elif table_data:
            ws.title = "데이터 현황"
            ws.append([f"분석 키워드: {query}", f"생성일시: {now_str}"])
            all_keys = []
            for row in table_data:
                for k in row.keys():
                    if k not in all_keys:
                        all_keys.append(k)
            ws.append(all_keys)
            for row in table_data:
                ws.append([row.get(k, "") for k in all_keys])
        else:
            ws.title = "리포트"
            ws.append([f"분석 키워드: {query}", f"생성일시: {now_str}"])
        buf = io.BytesIO()
        wb.save(buf)
        return buf.getvalue()


def _gen_excel_master(context_list, now_str) -> bytes:
    try:
        import xlsxwriter
        buf = io.BytesIO()
        wb = xlsxwriter.Workbook(buf)
        bold = wb.add_format({"bold": True})
        header_fmt = wb.add_format({"bold": True, "bg_color": "#1A237E", "font_color": "white"})

        # Summary sheet
        ws_sum = wb.add_worksheet("총괄 요약")
        ws_sum.write(0, 0, "전 분야 마스터 리포트", bold)
        ws_sum.write(1, 0, f"생성일시: {now_str}")
        ws_sum.write(3, 0, "분야", header_fmt)
        ws_sum.write(3, 1, "검색어", header_fmt)
        ws_sum.write(3, 2, "뉴스수", header_fmt)
        ws_sum.write(3, 3, "웹수", header_fmt)
        ws_sum.write(3, 4, "추세", header_fmt)
        ws_sum.write(3, 5, "변동률", header_fmt)

        for idx, item in enumerate(context_list):
            if not isinstance(item, dict):
                continue
            r = 4 + idx
            ws_sum.write(r, 0, item.get("expert", ""))
            ws_sum.write(r, 1, item.get("query", ""))
            ws_sum.write(r, 2, len(item.get("news", [])))
            ws_sum.write(r, 3, len(item.get("web", [])))
            trend = item.get("df", [])
            if trend and len(trend) >= 2:
                first_v = trend[0].get("Trend", 0)
                last_v = trend[-1].get("Trend", 0)
                if first_v > 0:
                    pct = ((last_v - first_v) / first_v) * 100
                    ws_sum.write(r, 4, "상승" if pct > 0 else "하락")
                    ws_sum.write(r, 5, round(pct, 1))

        ws_sum.set_column(0, 0, 15)
        ws_sum.set_column(1, 1, 40)

        # Per-expert sheets
        for i, item in enumerate(context_list):
            if not isinstance(item, dict):
                continue
            name = item.get("expert", f"분야{i + 1}")[:31]
            ws = wb.add_worksheet(name)
            ws.write(0, 0, f"분야: {name}", bold)
            ws.write(1, 0, f"검색어: {item.get('query', '')}")
            ws.write(2, 0, f"생성일시: {now_str}")
            ws.write(4, 0, "날짜", header_fmt)
            ws.write(4, 1, "트렌드", header_fmt)
            for j, row in enumerate(item.get("df", [])):
                ws.write(5 + j, 0, str(row.get("Date", "")))
                ws.write(5 + j, 1, row.get("Trend", 0))
            row_offset = 5 + len(item.get("df", [])) + 1
            ws.write(row_offset, 0, "뉴스 제목", header_fmt)
            ws.write(row_offset, 1, "출처", header_fmt)
            ws.write(row_offset, 2, "요약", header_fmt)
            ws.write(row_offset, 3, "링크", header_fmt)
            for j, n in enumerate(item.get("news", [])):
                _t = n.get("title", "")
                _s = n.get("snippet", "")
                if _s and _is_text_similar(_s, _t):
                    _s = ""
                ws.write(row_offset + 1 + j, 0, _t)
                ws.write(row_offset + 1 + j, 1, n.get("source", ""))
                ws.write(row_offset + 1 + j, 2, _s or "원문 참조")
                ws.write(row_offset + 1 + j, 3, n.get("link", ""))
            ws.set_column(0, 0, 30)
            ws.set_column(2, 2, 50)
            ws.set_column(3, 3, 60)
        wb.close()
        return buf.getvalue()
    except ImportError:
        return _gen_text_master(context_list, now_str)


# ═══════════════════════════════════════════════════════════════════════════════
# Main dispatcher
# ═══════════════════════════════════════════════════════════════════════════════

def _generate_local_report(report_format: str, context=None) -> bytes:
    """Generate report locally (standalone mode). context can be dict or list."""
    now_str = datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')

    # Master report: context is a list of dicts (one per expert domain)
    if isinstance(context, list):
        if report_format == "word":
            return _gen_word_master(context, now_str)
        elif report_format == "excel":
            return _gen_excel_master(context, now_str)
        else:
            return _gen_text_master(context, now_str)

    if isinstance(context, dict):
        query = context.get("query", "생활정보")
        news = context.get("news", [])
        web = context.get("web", [])
        youtube = context.get("youtube", [])
        raw_df = context.get("df", [])
        # Detect if df is trend data (has Trend key) or tabular snapshot data
        if raw_df and isinstance(raw_df[0], dict) and "Trend" in raw_df[0]:
            trend = raw_df
            table_data = []
        else:
            trend = []
            table_data = raw_df
    else:
        query = "생활정보"
        news = []
        web = []
        youtube = []
        trend = []
        table_data = []

    if report_format == "excel":
        return _gen_excel(query, news, web, trend, now_str, table_data=table_data, youtube=youtube)
    elif report_format == "word":
        return _gen_word(query, news, web, trend, now_str, table_data=table_data, youtube=youtube)
    else:
        return _gen_text(query, news, web, trend, now_str, table_data=table_data, youtube=youtube)


def download_report_from_api(report_format: str, context: dict = None):
    """Call the backend API to generate a report."""
    url = f"{API_BASE_URL}/report/generate"
    payload = {"report_type": report_format}
    if context:
        payload["context"] = context
    try:
        response = requests.post(url, json=payload, stream=True, timeout=10)
        response.raise_for_status()
        return response.content
    except Exception:
        return None


def render_download_buttons(context: dict = None):
    """Render report download section. Works in both API and standalone mode.

    Pre-generates all formats on first render to avoid page scroll jump
    when download buttons appear dynamically.
    """
    if context:
        if isinstance(context, list):
            st.markdown("### 📥 전 분야 마스터 리포트 다운로드")
            st.caption("전문가 최신 동향 데이터를 총망라한 마스터 리포트를 다운로드합니다.")
        else:
            st.markdown("### 📥 전문가 맞춤형 보고서 다운로드")
            st.caption(f"'{context.get('query', '분석 키워드')}'에 대한 최신 동향과 뉴스를 포함한 심층 리포트를 다운로드합니다.")
    else:
        st.markdown("### 📥 통합 보고서 다운로드")
        st.caption("현재 수집된 날씨, 뉴스, 교통 요약 리포트를 다운로드합니다.")

    # Compute a stable cache key from context to avoid regeneration on every rerun
    import hashlib
    ctx_hash = ""
    if context:
        try:
            ctx_str = str(context.get("query", "")) if isinstance(context, dict) else str(len(context))
            ctx_hash = hashlib.md5(ctx_str.encode()).hexdigest()[:8]
        except Exception:
            ctx_hash = "default"

    now_str = datetime.datetime.now().strftime('%Y%m%d_%H%M')
    prefix = "Master_Report" if isinstance(context, list) else ("Expert_Report" if context else "LifeInfo_Summary")

    formats = [
        ("Word (.docx)", "word", "application/vnd.openxmlformats-officedocument.wordprocessingml.document", "📝", ".docx"),
        ("텍스트 (.txt)", "text", "text/plain", "📄", ".txt"),
        ("Excel (.xlsx)", "excel", "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", "📊", ".xlsx"),
    ]

    # Use session_state to cache generated reports, avoiding regeneration + scroll jump
    cache_key = f"_report_cache_{ctx_hash}"
    if cache_key not in st.session_state:
        st.session_state[cache_key] = {}

    cached = st.session_state[cache_key]

    cols = st.columns(3)
    for i, (label, fmt, mime_type, icon, ext) in enumerate(formats):
        with cols[i]:
            filename = f"{prefix}_{now_str}{ext}"

            if fmt in cached and cached[fmt]:
                # Already generated — show download button directly (no layout shift)
                st.download_button(
                    label=f"{icon} {label} 저장",
                    data=cached[fmt],
                    file_name=filename,
                    mime=mime_type,
                    key=f"dl_{fmt}_{ctx_hash}",
                )
            else:
                # Show generate button
                if st.button(f"{icon} {label} 생성", key=f"btn_gen_{fmt}_{ctx_hash}"):
                    with st.spinner(f"{label} 생성 중..."):
                        content = None
                        if IS_API_MODE:
                            content = download_report_from_api(fmt, context)
                        if not content:
                            content = _generate_local_report(fmt, context)
                        if content:
                            cached[fmt] = content
                            st.session_state[cache_key] = cached
                            st.rerun()
